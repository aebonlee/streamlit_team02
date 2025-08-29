import os, io, re, json, textwrap, datetime
from typing import Optional, List, Dict, Tuple

import streamlit as st

# ===== Optional libraries =====
# pandas, numpy: 데이터프레임 및 수치 계산용
try:
    import pandas as pd
    import numpy as np
    PANDAS_OK = True
except Exception:
    PANDAS_OK = False  # 설치 안 되어 있을 때 False 처리

# 시각화 라이브러리
try:
    import altair as alt
    import plotly.express as px
    VIZ_OK = True
except Exception:
    VIZ_OK = False

# 웹 요청/크롤링
try:
    import requests
    from bs4 import BeautifulSoup
    HTTP_OK = True
except Exception:
    HTTP_OK = False

# docx 파일 읽기
try:
    from docx import Document
    DOCX_OK = True
except Exception:
    DOCX_OK = False

# PDF 생성
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph
    from reportlab.lib.styles import getSampleStyleSheet
    REPORTLAB_OK = True
except Exception:
    REPORTLAB_OK = False

# LLM (OpenAI) 관련
try:
    from langchain_openai import ChatOpenAI
    from langchain.prompts import ChatPromptTemplate
    from langchain.chains import LLMChain
    LLM_OK = True
except Exception:
    LLM_OK = False

# ================= 세션 초기화 =================
# Streamlit 세션 상태 초기화
if "messages" not in st.session_state:
    # AI 첫 인사 메시지
    st.session_state.messages = [
        {
            "role": "ai",
            "content": "안녕하세요! AI 자기소개서 코치입니다. 무엇을 도와드릴까요?",
            "time": datetime.datetime.now().strftime("%H:%M"),
        }
    ]
if "saved_files" not in st.session_state:
    st.session_state.saved_files = []  # 업로드된 파일 리스트
if "api_key" not in st.session_state:
    st.session_state.api_key = os.getenv("OPENAI_API_KEY", "")  # 환경변수 OpenAI 키
if "basic_settings" not in st.session_state:
    st.session_state.basic_settings = {
        "model": "GPT-4 (무료)",  # 기본 모델
        "tone": "전문적",        # 문체
        "length": 800,           # 출력 길이
    }
if "advanced_settings" not in st.session_state:
    st.session_state.advanced_settings = {
        "creativity": 0.5,        # 창의성 정도
        "export_format": "PDF 문서",  # 내보내기 형식
    }
if "show_guide" not in st.session_state:
    st.session_state.show_guide = False  # 가이드 표시 여부

# ================= 가이드라인 =================
GUIDE = """📝 **AI 자기소개서 코치 사용 가이드**
1) **자소서 평가 탭**에서 텍스트를 붙여넣고, 회사/직무를 입력 후 **평가 실행**
   - 규칙 기반 점수 + LLM 개선안 + 스킬 매칭 표 제공
2) **트렌드/기업 탭**에서 회사/직무 입력 → 최신 공고/기술 추이와 회사 인재상 요약
3) 좌측 **설정**에서 OpenAI 및 (선택) SERP/Bing 키를 입력하면 웹 요약 기능 활성화
4) (선택) Tableau Public 링크가 있다면 탭 하단에 임베드하여 팀과 공유 가능

📡 **새 기능 안내**
채팅창에 **"내가 (회사명)의 자소서에 대한 데이터를 얻고 싶어"** 라고 입력하면,
로컬 CSV를 바탕으로 해당 기업의 **채용/기술 수요 요약**을 바로 알려드립니다!
"""

def get_guideline() -> str:
    return GUIDE  # 가이드 텍스트 반환

# ================= 데이터 경로/헬퍼 =================
def _env(key: str, default: str = "") -> str:
    return os.getenv(key, default)  # 환경변수 읽기

@st.cache_data(show_spinner=False)
def _default_data_dir() -> str:
    # 기본 데이터 디렉토리 설정
    if os.path.isdir("/mnt/data"):
        return "/mnt/data"
    return _env("DATA_DIR", "./data")

DATA_DIR = _default_data_dir()  # 데이터 폴더 경로

@st.cache_data(show_spinner=False)
def load_csv(name: str) -> Optional[pd.DataFrame]:
    # CSV 파일 불러오기
    if not PANDAS_OK:
        return None
    candidates = [os.path.join(DATA_DIR, name), os.path.join(".", name)]
    for path in candidates:
        if os.path.isfile(path):
            try:
                return pd.read_csv(path)
            except Exception:
                pass
    return None

# 데이터 로드
if PANDAS_OK:
    job_market = load_csv("job_market.csv")          # 채용시장 데이터
    macro = load_csv("macro_indicators.csv")        # 거시 지표
    skills = load_csv("skills_analysis.csv")        # 스킬 분석
    tech_trends = load_csv("tech_trends.csv")       # 기술 트렌드
else:
    job_market = macro = skills = tech_trends = None

# ================= 텍스트/문서 처리 =================
def read_text_from_upload(uploaded) -> str:
    # 업로드된 파일에서 텍스트 읽기
    if uploaded is None:
        return ""
    name = uploaded.name.lower()
    try:
        if name.endswith(".txt"):
            return uploaded.read().decode("utf-8", errors="ignore")
        if name.endswith(".docx") and DOCX_OK:
            doc = Document(uploaded)
            return "\n".join(p.text for p in doc.paragraphs)
        return uploaded.read().decode("utf-8", errors="ignore")
    except Exception as e:
        return f"[파일 읽기 오류] {e}"

# ================= 규칙 기반 스코어러 =================
# 행동 동사 목록: 성과/능력 표현
ACTION_WORDS = [
    "개선", "최적화", "설계", "구현", "분석", "자동화", "협업", "리팩터", "검증",
    "성과", "증가", "감소", "달성", "기여", "해결", "리더", "조율",
]
# STAR 기법 토큰
STAR_TOKENS = ["상황", "과제", "행동", "결과", "Situation", "Task", "Action", "Result"]
# 불필요하게 반복되는 단어
FILLERS = ["최대한", "정말", "매우", "다양한", "많은", "열정", "성실", "노력"]
# 숫자/백분율 정규표현
NUM_RE = re.compile(r"(?<!\w)(?:[0-9]+(?:\.[0-9]+)?%?|[일이삼사오육칠팔구십]+%?)(?!\w)")

def tokenize_kr(text: str) -> List[str]:
    # 한글/영문 단어 단위 토큰화
    return re.findall(r"[\w가-힣%]+", text.lower())

def skill_coverage(text: str, skills_df: Optional[pd.DataFrame], month: Optional[str] = None) -> Tuple[float, List[str]]:
    # 스킬 커버리지 계산
    if skills_df is None or len(skills_df) == 0:
        return 0.0, []
    toks = set(tokenize_kr(text))
    df = skills_df.copy()
    if month and "month" in df.columns:
        df = df[df["month"] == month] if (df["month"] == month).any() else df
    listed = [str(s).lower() for s in df["skill"].unique().tolist()]
    matched = sorted({s for s in listed if any(s in t for t in toks)})
    cov = len(matched) / max(1, len(set(listed)))
    return cov, matched[:20]

def compute_resume_scores(text: str, role: str = "", company: str = "", skills_df: Optional[pd.DataFrame] = None) -> Dict[str, float]:
    # 자기소개서 점수 계산
    tokens = tokenize_kr(text)
    n_words = len(tokens)
    n_chars = len(text)
    nums = NUM_RE.findall(text)  # 숫자 등장 빈도
    metric_density = min(1.0, len(nums) / max(1, n_words) * 10)  # 숫자 밀도 스코어
    action_hits = sum(1 for w in ACTION_WORDS if any(w in t for t in tokens))  # 행동 동사 사용 횟수
    action_score = min(1.0, action_hits / 6)
    star_hits = sum(1 for w in STAR_TOKENS if any(w.lower() in t for t in tokens))  # STAR 기법 사용
    star_score = min(1.0, star_hits / 4)
    filler_hits = sum(tokens.count(f.lower()) for f in FILLERS)  # 불필요 단어 카운트
    filler_penalty = min(0.3, filler_hits / max(1, n_words) * 5)  # 페널티
    length_score = 1.0 if 600 <= n_chars <= 1200 else max(0.3, 1 - abs(n_chars - 900) / 1200)  # 길이 적절성
    month = None
    if skills_df is not None and "month" in skills_df.columns:
        month = skills_df["month"].max()
    cov, matched = skill_coverage(text, skills_df, month)  # 스킬 커버리지
    coverage_score = min(1.0, 0.5 + cov)
    # 가중치 설정
    weights = {
        "metrics": 0.25,
        "action": 0.15,
        "star": 0.15,
        "length": 0.15,
        "coverage": 0.30,
    }
    # 종합 점수 계산
    total = (
        metric_density * weights["metrics"]
        + action_score * weights["action"]
        + star_score * weights["star"]
        + length_score * weights["length"]
        + coverage_score * weights["coverage"]
    )
    total = max(0.0, min(1.0, total - filler_penalty))  # 페널티 적용
    return {
        "총점(0-100)": round(total * 100, 1),
        "성과(숫자)밀도": round(metric_density, 3),   # <- 여기까지 주석 완료
    }
"행동성": round(action_score, 3),  # ACTION_WORDS 기반 점수, 자기소개서에서 구체적 행동을 얼마나 강조했는지
"STAR구조": round(star_score, 3),  # STAR 구조(Situation, Task, Action, Result) 준수 정도
"길이적정": round(length_score, 3),  # 글 길이가 적절한지 점수화
"스킬커버리지": round(coverage_score, 3),  # 스킬 매칭 커버리지 점수
"군더더기감점": round(filler_penalty, 3),  # 군더더기/불필요한 단어 사용에 따른 감점
}

def llm_improve(text: str, role: str, company: str, tone: str, length: int) -> str:
    # LLM(대규모 언어 모델) 기반 자기소개서 개선 함수
    if not LLM_OK or not st.session_state.get("api_key"):
        # LLM 사용 불가 시 안내 메시지 반환
        return "[LLM 미사용] OpenAI API 키가 없거나 라이브러리가 없습니다. 설정 탭에서 API 키를 입력하세요."
    
    # 시스템 프롬프트 설정: 톤, 최대 길이, STAR 구조 적용 등
    system = f"""당신은 한국어 자기소개서 첨삭 전문가입니다.
    - 톤: {tone}
    - 최대 길이: {length}자
    - 작업: 아래 자기소개서를 {company} {role} 지원 기준으로 STAR 구조와 수치 중심으로 다듬고, 중복/군더더기를 줄이세요.
    - 출력 형식: 1) 개선 요약(불릿) 2) 개선된 자기소개서(문단) 3) 다음 액션 3가지"""
    
    # 프롬프트 템플릿 생성
    tmpl = ChatPromptTemplate.from_messages([("system", system), ("human", "원문:\n{orig}")])
    
    # LLM 체인 구성
    chain = LLMChain(
        llm=ChatOpenAI(
            api_key=st.session_state.get("api_key", os.getenv("OPENAI_API_KEY")),
            model="gpt-4o-mini",  # 사용할 모델
            temperature=0.4,      # 창의성 정도
        ),
        prompt=tmpl,
    )
    
    # LLM 호출
    out = chain.invoke({"orig": text})
    return out.get("text", str(out))  # 결과 반환

# ================= 채팅용 기능 =================
COMPANY_CMD_RE = re.compile(
    r"내가\s*(?P<company>.+?)\s*의?\s*자소서에\s*대한\s*데이터(?:를)?\s*얻고\s*싶어",
    re.IGNORECASE,
)  # 사용자 입력에서 회사명 추출 정규식

def _clean_company(s: str) -> str:
    # 회사명 문자열 앞뒤 불필요한 문자 제거
    s = s.strip()
    s = re.sub(r'^[\"“”‘’\'(\[]+', "", s)
    s = re.sub(r'[\"“”‘’\'\])]+$', "", s)
    return s.strip()

def try_parse_company_query(text: str) -> Optional[str]:
    # 사용자 문장에서 회사명 파싱
    if not text:
        return None
    m = COMPANY_CMD_RE.search(text)
    if not m:
        return None
    return _clean_company(m.group("company"))

def summarize_company_from_csvs(company: str) -> str:
    # 로컬 CSV 데이터를 기반으로 회사 자소서 요약
    if not PANDAS_OK:
        return (
            f"### 📊 기업 자소서 데이터 요약 — {company}\n"
            "- 이 기능을 사용하려면 `pandas` 설치가 필요해요. `pip install pandas` 후 다시 시도해주세요.\n"
        )
    lines = [f"### 📊 기업 자소서 데이터 요약 — {company}"]
    
    # job_market.csv 기반 공고 수 확인
    if job_market is not None:
        sub = job_market.copy()
        if "company" in sub.columns:
            sub = sub[sub["company"].astype(str).str.contains(company, case=False, na=False)]
        try:
            cnt = sub["job_code"].nunique() if "job_code" in sub.columns else len(sub)
        except Exception:
            cnt = len(sub)
        recent = ""
        if "posted_date" in sub.columns:
            try:
                _d = pd.to_datetime(sub["posted_date"], errors="coerce")
                if _d.notna().any():
                    recent = _d.max().date().isoformat()
            except Exception:
                pass
        msg = f"- 최근 수집 공고 수: **{cnt}건**"
        if recent:
            msg += f" (최신: {recent})"
        lines.append(msg)
    else:
        lines.append("- `job_market.csv`를 찾지 못했습니다. `/mnt/data` 또는 프로젝트 루트에 배치해주세요.")
    
    # skills_analysis.csv 기반 상위 기술 수요
    if skills is not None and "skill" in skills.columns:
        kdf = skills.copy()
        if "month" in kdf.columns and kdf["month"].notna().any():
            top_month = kdf["month"].max()
            if (kdf["month"] == top_month).any():
                kdf = kdf[kdf["month"] == top_month]
        try:
            if "job_count" in kdf.columns:
                top_skills = (
                    kdf.groupby("skill")["job_count"].sum().sort_values(ascending=False).head(10).index.tolist()
                )
            else:
                top_skills = kdf["skill"].value_counts().head(10).index.tolist()
        except Exception:
            top_skills = []
        if top_skills:
            lines.append(f"- 최근 상위 기술 수요: {', '.join(top_skills)}")
    else:
        lines.append("- `skills_analysis.csv`를 찾지 못해 상위 기술 수요를 계산할 수 없습니다.")
    
    return "\n".join(lines) + "\n\n> *참고: 데이터는 로컬 CSV 기준 요약이며, 더 자세한 웹 리서치는 선택적으로 확장 가능합니다.*"

# ===== 웹 동향/기업 인재상 수집 =====
def search_web(query: str, topk: int = 5) -> List[Dict[str, str]]:
    # Google SERP 또는 Bing API를 통해 웹 검색
    res: List[Dict[str, str]] = []
    serp_key = os.getenv("SERPAPI_API_KEY")
    bing_key = os.getenv("BING_API_KEY")
    try:
        if serp_key:
            params = {"engine": "google", "q": query, "api_key": serp_key, "num": topk, "hl": "ko"}
            r = requests.get("https://serpapi.com/search.json", params=params, timeout=15)
            j = r.json()
            for it in j.get("organic_results", [])[:topk]:
                res.append({"title": it.get("title", ""), "url": it.get("link", ""), "snippet": it.get("snippet", "")})
        elif bing_key:
            headers = {"Ocp-Apim-Subscription-Key": bing_key}
            r = requests.get(
                "https://api.bing.microsoft.com/v7.0/search",
                params={"q": query, "count": topk, "mkt": "ko-KR"},
                headers=headers,
                timeout=15,
            )
            j = r.json()
            for it in j.get("webPages", {}).get("value", [])[:topk]:
                res.append({"title": it.get("name", ""), "url": it.get("url", ""), "snippet": it.get("snippet", "")})
        return res
    except Exception:
        return res

def fetch_and_summarize(urls: List[str]) -> str:
    # URL 리스트에서 텍스트 수집 후 요약
    texts = []
    for u in urls[:5]:
        try:
            html = requests.get(u, timeout=15).text
            soup = BeautifulSoup(html, "html.parser")
            t = " ".join([p.get_text(" ", strip=True) for p in soup.find_all(["p", "li"])])
            texts.append(textwrap.shorten(t, 3000))
        except Exception:
            pass
    joined = "\n\n".join(texts) if texts else ""
    if not joined:
        return "(웹 페이지에서 요약할 텍스트를 수집하지 못했습니다.)"
    
    # LLM 사용 가능 시 요약 생성
    if LLM_OK and os.getenv("OPENAI_API_KEY"):
        sys = "너는 리서치 요약가다. 한국어로 5개 불릿, 5줄 이하 요약으로 정리하라."
        tmpl = ChatPromptTemplate.from_messages([("system", sys), ("human", "다음 자료를 요약:\n{t}")])
        out = LLMChain(llm=ChatOpenAI(model="gpt-4o-mini", temperature=0.2), prompt=tmpl).invoke({"t": joined})
        return out.get("text", str(out))
    return joined[:1500]

def company_persona_and_requirements(company: str, role: str) -> Dict[str, str]:
    # 회사 인재상 및 요구역량 요약
    result = {"인재상": "", "요구역량": "", "출처": []}
    if not HTTP_OK:
        return result
    queries = [
        f"{company} 인재상 site:co.kr OR site:com OR site:kr",
        f"{company} 채용 {role} 자기소개서",
        f"{company} core values culture",
    ]
    urls = []
    for q in queries:
        hits = search_web(q, topk=5)
        urls.extend([h["url"] for h in hits])
        result["출처"].extend(hits)
    urls = list(dict.fromkeys([u for u in urls if u]))  # 중복 제거
    summary = fetch_and_summarize(urls)
    result["인재상"] = summary
    
    # 스킬 데이터가 있으면 요구역량 예시 추가
    if skills is not None and "skill" in skills.columns:
        top_month = skills["month"].max() if "month" in skills.columns else None
        kdf = skills.copy()
        if top_month:
            kdf = kdf[kdf["month"] == top_month]
        top_skills = (
            kdf.groupby("skill")["job_count"].sum().sort_values(ascending=False).head(10).index.tolist()
            if "job_count" in kdf.columns
            else kdf["skill"].value_counts().head(10).index.tolist()
        )
        result["요구역량"] = "최근 수요 상위 기술 예시: " + ", ".join(top_skills)
    return result

# ================= AI 응답 생성 =================
def get_ai_response(user_input: str, uploaded_file=None) -> str:
    # 가이드 키워드 처리
    guideline_keywords = ["가이드", "가이드라인", "도움말", "사용법", "어떻게"]
    if any(keyword in user_input for keyword in guideline_keywords):
        return get_guideline()
    
    # 기본 템플릿
    if not st.session_state.basic_settings or not LLM_OK or not st.session_state.get("api_key"):
        templates = {
            "default": """자기소개서 작성을 도와드리겠습니다!
구체적으로 알려주시면 더 정확한 도움을 드릴 수 있어요:
• 어떤 직무에 지원하시나요?
• 어떤 부분이 어려우신가요?
• 특별히 강조하고 싶은 경험이 있나요?""",
            "첨삭": """자기소개서 첨삭 포인트를 알려드릴게요:
✅ 구체적인 숫자와 성과 포함
✅ 직무와 연관된 경험 강조
✅ 문장은 간결하고 명확하게
✅ 진정성 있는 지원동기
파일을 업로드하거나 내용을 보내주시면 더 자세히 봐드릴게요!""",
            "시작": """자기소개서 작성을 시작해볼까요?
**Step 1. 기본 정보**
• 지원 회사:  # 사용자가 지원할 회사명 입력
• 지원 직무:  # 사용자가 지원할 직무 입력
• 경력 구분: (신입/경력)  # 지원자의 경력 여부 선택"""
        }
"""
사용자 입력 기반 AI 자기소개서 코칭 함수 및 대화 UI 관련
"""

# ================= AI 응답 생성 (계속) =================
# 사용자가 도움말/가이드 요청 시 미리 정의된 템플릿 반환
guideline_keywords = ["가이드", "가이드라인", "도움말", "사용법", "어떻게"]
if any(keyword in user_input for keyword in guideline_keywords):
    return get_guideline()

# 기본 템플릿 정의
templates = {
    "default": """이 정보를 알려주시면 맞춤형으로 도와드릴게요!""",
    "첨삭": """자기소개서 첨삭 포인트:
✅ 구체적 숫자/성과 포함
✅ 직무 관련 경험 강조
✅ 문장 간결/명확
✅ 진정성 있는 지원동기""",
    "시작": """자기소개서 작성 시작 안내:
• 지원 회사:
• 지원 직무:
• 경력 구분: (신입/경력)""",
    "예시": """자기소개서 예시:
"문제 해결 능력을 바탕으로 한 프로젝트 경험을 통해 팀에 기여했던 사례가 있습니다."
→ 경험을 구체적으로 설명하는 방식"""
}

# 사용자가 입력한 키워드에 따라 템플릿 반환
if "첨삭" in user_input or "수정" in user_input:
    return templates["첨삭"]
elif "시작" in user_input or "처음" in user_input:
    return templates["시작"]
elif "예시" in user_input:
    return templates["예시"]
else:
    return templates["default"]

# ================= LLM 기반 일반 응답 =================
try:
    # 선택 모델 매핑
    model_map = {
        "GPT-4 (무료)": "gpt-4o-mini",
        "GPT-4": "gpt-4o",
        "GPT-3.5": "gpt-3.5-turbo",
    }
    selected_model = st.session_state.basic_settings.get("model", "GPT-4 (무료)")
    model_name = model_map.get(selected_model, "gpt-4o-mini")

    # LLM 객체 생성
    llm = ChatOpenAI(
        api_key=st.session_state.get("api_key"),
        model=model_name,
        temperature=st.session_state.advanced_settings.get("creativity", 0.5),
    )

    # 시스템 프롬프트 설정: 전문 자기소개서 코치 톤
    system_prompt = f"""당신은 전문 자기소개서 작성 코치입니다.
    톤: {st.session_state.basic_settings['tone']}
    최대 길이: {st.session_state.basic_settings['length']}자

    - 구체적이고 실용적인 조언
    - 예시 포함
    - 친근하면서 전문적인 톤
    - 이모지 최소 사용"""

    # 파일 업로드 처리(txt, docx)
    if uploaded_file:
        try:
            if uploaded_file.name.endswith('.txt'):
                content = uploaded_file.read().decode('utf-8')
            elif uploaded_file.name.endswith('.docx') and DOCX_OK:
                doc = Document(uploaded_file)
                content = '\n'.join([p.text for p in doc.paragraphs])
            else:
                content = "파일을 읽을 수 없습니다."
            user_input = f"다음 자기소개서를 검토하고 개선점을 제안해주세요:\n\n{content}\n\n{user_input}"
        except Exception as e:
            return f"파일 처리 중 오류: {e}"

    # 프롬프트 템플릿 생성
    prompt = ChatPromptTemplate.from_messages([("system", system_prompt), ("human", "{input}")])
    chain = LLMChain(llm=llm, prompt=prompt)
    response = chain.invoke({"input": user_input})
    return response.get("text", str(response))
except Exception as e:
    return f"오류가 발생했습니다. 다시 시도해주세요.\n{str(e)}"

# ================= 대화 저장 =================
def save_conversation() -> str:
    # 현재 세션 메시지를 저장
    content = ""
    for msg in st.session_state.messages:
        role = "👤 사용자" if msg["role"] == "user" else "🤖 AI 코치"
        content += f"[{msg.get('time', '')}] {role}\n{msg['content']}\n\n"

    # 파일명 및 형식 설정
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"자소서대화_{timestamp}"
    export = st.session_state.advanced_settings.get("export_format", "텍스트 파일")

    # PDF 저장
    if export == "PDF 문서" and REPORTLAB_OK:
        bio = io.BytesIO()
        doc = SimpleDocTemplate(bio, pagesize=letter)
        styles = getSampleStyleSheet()
        story = [Paragraph(p, styles["Normal"]) for p in content.split('\n')]
        doc.build(story)
        file_data = bio.getvalue()
        mime = "application/pdf"
        ext = "pdf"

    # Word(docx) 저장
    elif export == "Word 문서" and DOCX_OK:
        doc = Document()
        doc.add_heading('AI 자기소개서 코칭 대화', 0)
        for para in content.split('\n'):
            doc.add_paragraph(para)
        bio = io.BytesIO()
        doc.save(bio)
        file_data = bio.getvalue()
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext = "docx"

    # HTML 저장
    elif export == "HTML 문서":
        file_data = f"<html><body><pre>{content}</pre></body></html>".encode("utf-8")
        mime = "text/html"
        ext = "html"

    # 일반 텍스트 저장
    else:
        file_data = content.encode("utf-8")
        mime = "text/plain"
        ext = "txt"

    # 세션에 저장
    st.session_state.saved_files.append(
        {
            "name": f"{filename}.{ext}",
            "date": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
            "size": len(file_data),
            "data": file_data,
            "mime": mime,
        }
    )
    return f"{filename}.{ext}"

# ================= 페이지 설정 및 스타일 =================
st.set_page_config(page_title="AI 자기소개서 코칭+", page_icon="💬", layout="wide")
MAIN = "#22C55E"  # 메인 색
BG = "#F5FBFB"    # 배경색
USER_BG = "#DCFCE7"  # 사용자 채팅 배경
BOT_BG = "#F3F4F6"   # AI 채팅 배경

# 사용자/봇 UI 스타일 정의
st.markdown(
    f"""
    <style>
        body {{ background:{BG}; }}
        .header {{position:sticky; top:0; background:{MAIN}; color:white; padding:34px; margin-top:0; margin-bottom:14px; border-radius:10px; z-index:10; font-size:4em; font-weight:900; text-align:center;}}
        .bubble-user {{background:{USER_BG}; padding:10px; border-radius:16px; margin:6px 0; text-align:right}}
        .bubble-bot {{background:{BOT_BG}; padding:10px; border-radius:16px; margin:6px 0; text-align:left}}
        .metric-box {{border-radius:14px; padding:10px; background:white; border:1px solid #e5e7eb}}
    </style>
    """,
    unsafe_allow_html=True,
)

# ================= 사이드바 설정 =================
with st.sidebar:
    st.header("설정")
    api_key = st.text_input("OpenAI API Key", value=os.getenv("OPENAI_API_KEY", ""), type="password")
    if api_key:
        os.environ["OPENAI_API_KEY"] = api_key
        st.session_state.api_key = api_key

    serp_key = st.text_input("SERPAPI_API_KEY (선택)", value=os.getenv("SERPAPI_API_KEY", ""), type="password")
    if serp_key:
        os.environ["SERPAPI_API_KEY"] = serp_key

    bing_key = st.text_input("BING_API_KEY (선택)", value=os.getenv("BING_API_KEY", ""), type="password")
    if bing_key:
        os.environ["BING_API_KEY"] = bing_key

    st.markdown("---")
    st.subheader("모델/톤")
    models = ["GPT-4 (무료)", "GPT-4", "GPT-3.5"]
    st.session_state.basic_settings["model"] = st.selectbox(
        "AI 모델", models, index=models.index(st.session_state.basic_settings.get("model", models[0]))
    )
    tones = ["전문적", "친근한", "격식 있는", "캐주얼"]
    st.session_state.basic_settings["tone"] = st.selectbox(
        "작성 톤", tones, index=tones.index(st.session_state.basic_settings.get("tone", tones[0]))
    )
    st.session_state.basic_settings["length"] = st.slider(
        "글자 수", min_value=300, max_value=2000, value=st.session_state.basic_settings.get("length", 800)
    )

    st.markdown("---")
    st.subheader("세부 설정")
    st.session_state.advanced_settings["creativity"] = st.slider(
        "창의성", 0.0, 1.0, value=st.session_state.advanced_settings.get("creativity", 0.5)
    )
    export_options = ["PDF 문서", "Word 문서", "텍스트 파일", "HTML 문서"]
    st.session_state.advanced_settings["export_format"] = st.selectbox(
        "내보내기 형식", export_options, index=export_options.index(st.session_state.advanced_settings.get("export_format", "PDF 문서"))
    )

# 페이지 상단 헤더
st.markdown(f"<div class='header'><b>AI 자기소개서 코칭 +</b></div>", unsafe_allow_html=True)

# ================= 탭 UI =================
tab_chat, tab_eval, tab_trend = st.tabs(["💬 대화", "🧭 자소서 평가", "📈 트렌드/기업"])

# --------- 💬 대화 ---------
with tab_chat:
    col_title, col_button = st.columns([3, 1])
    with col_title:
        st.subheader("일반 코칭 대화")
    with col_button:
        if st.button("📖 사용 가이드", use_container_width=True):
            st.session_state.show_guide = not st.session_state.show_guide

    # 가이드 표시
    if st.session_state.show_guide:
        st.markdown(
            "<div style='background-color: #f0f2f6; padding: 15px; border-radius: 15px; margin: 10px 0; border-left: 4px solid #22C55E;'>",
            unsafe_allow_html=True,
        )
        st.markdown(GUIDE, unsafe_allow_html=True)
        st.markdown("</div>", unsafe_allow_html=True)

    # 파일 업로드
    uploaded_file = st.file_uploader("📎 파일 첨부 (txt, docx)", type=["txt", "docx"], key="chat_file")

    # 이전 메시지 출력
    for msg in st.session_state.messages:
        if msg["role"] == "user":
            st.markdown(f"<div class='bubble-user'>{msg['content']}</div>", unsafe_allow_html=True)
        else:
            content_html = msg['content'].replace('\n', '<br>')
            st.markdown(f"<div class='bubble-bot'>{content_html}</div>", unsafe_allow_html=True)

    # 사용자 입력
    user_input = st.text_input("메시지 입력", key="chat_input")
    col_send, col_save = st.columns(2)
    if col_send.button("전송", use_container_width=True):
        if user_input.strip():
            # 세션에 사용자 메시지 저장
            st.session_state.messages.append(
                {"role": "user", "content": user_input.strip(), "time": datetime.datetime.now().strftime("%H:%M")}
            )
            # 회사명 파싱 후 요약 또는 일반 응답
            company_name = try_parse_company_query(user_input)
            if company_name:
                response = summarize_company_from_csvs(company_name)
            else:
                response = get_ai_response(user_input, uploaded_file)
            # 세션에 AI 응답 저장
            st.session_state.messages.append(
                {"role": "ai", "content": response, "time": datetime.datetime.now().strftime("%H:%M")}
            )
# ================= 대화 저장 & 다운로드 =================
# 사용자가 '전송' 후 대화 내용을 파일로 저장하거나 다운로드할 수 있는 기능 구현
if col_save.button("대화 저장", use_container_width=True):
    # save_conversation 함수 호출 → 대화 내용 파일로 저장
    fname = save_conversation()
    st.success(f"{fname} 저장됨! 아래에서 다운로드하세요.")  # 저장 완료 메시지 출력

# 이미 저장된 파일이 있으면 다운로드 가능
if st.session_state.saved_files:
    with st.expander("📂 저장된 파일"):
        for i, file in enumerate(st.session_state.saved_files):
            st.write(f"📄 {file['name']} ({file['date']}, {file['size']} bytes)")
            st.download_button(
                label="다운로드",
                data=file["data"],          # 파일 데이터
                file_name=file["name"],     # 파일 이름
                mime=file["mime"],          # MIME 타입
                key=f"download_{i}",        # 다운로드 버튼 고유 키
            )

# ================= 🧭 자소서 평가 탭 =================
with tab_eval:
    st.subheader("자소서 평가")  # 탭 제목
    colL, colR = st.columns([1, 1])  # 좌/우 2 컬럼 레이아웃

    # ----- 왼쪽 컬럼: 입력 영역 -----
    with colL:
        text = st.text_area("자기소개서 텍스트")          # 평가할 자기소개서 입력
        company = st.text_input("지원 회사")             # 지원 회사 입력
        role = st.text_input("지원 직무")               # 지원 직무 입력
        tone = st.selectbox("톤", ["전문적", "친근한", "격식 있는", "캐주얼"], index=0)
        length = st.slider("출력 길이(자)", 400, 1500, 900)  # 개선안 출력 길이 설정
        run = st.button("평가 실행", type="primary")   # 평가 실행 버튼

    # ----- 오른쪽 컬럼: 출력 영역 -----
    with colR:
        st.markdown("**평가 지표**")
        st.caption("성과·행동·STAR·길이·스킬커버리지·군더더기")  # 참고 지표
        placeholder_metrics = st.empty()  # 평가 점수 표시용 placeholder
        improved_box = st.empty()         # 개선안 표시용 placeholder

    # 평가 실행 버튼 클릭 시
    if run and text.strip():
        with st.spinner("평가 중…"):
            scores = compute_resume_scores(text, role, company, skills)  # 점수 계산

        # 평가 점수 시각화
        with placeholder_metrics.container():
            if VIZ_OK and PANDAS_OK:
                df_score = pd.DataFrame(
                    {
                        "항목": ["총점", "성과", "행동", "STAR", "길이", "스킬", "감점"],
                        "점수": [
                            scores['총점(0-100)'],
                            scores['성과(숫자)밀도']*100,
                            scores['행동성']*100,
                            scores['STAR구조']*100,
                            scores['길이적정']*100,
                            scores['스킬커버리지']*100,
                            scores['군더더기감점']*100,
                        ],
                    }
                )
                # Plotly 막대 그래프 시각화
                fig = px.bar(df_score, x="항목", y="점수", title="평가 결과(%)", range_y=[0, 100])
                st.plotly_chart(fig, use_container_width=True)
            st.json(scores)  # 점수 JSON 형태로 출력

        # 개선안 작성
        with st.spinner("개선안 작성…"):
            improved = llm_improve(text, role, company, tone, length)  # LLM 기반 개선안 생성
        with improved_box.container():
            st.markdown("### ✍️ 개선안")
            st.markdown(improved)

        # 스킬 커버리지 분석
        if PANDAS_OK and skills is not None:
            cov, matched = skill_coverage(text, skills)
            st.markdown("---")
            st.markdown("**스킬 매칭(최근 수요 기준)**")
            st.write(f"커버리지: {cov*100:.1f}% / 매칭: {', '.join(matched) if matched else '(없음)'}")
            # 상위 기술 수요 차트
            if VIZ_OK:
                top_month = skills['month'].max() if 'month' in skills.columns else None
                kdf = skills.copy()
                if top_month:
                    kdf = kdf[kdf['month']==top_month]
                if 'job_count' in kdf.columns:
                    kdf = kdf.groupby('skill')['job_count'].sum().sort_values(ascending=False).head(15).reset_index()
                    st.altair_chart(
                        alt.Chart(kdf).mark_bar().encode(
                            x='job_count', 
                            y=alt.Y('skill', sort='-x')
                        ).properties(height=380, title=f"{top_month or ''} 상위 기술 수요"),
                        use_container_width=True,
                    )

# ================= 📈 트렌드/기업 탭 =================
with tab_trend:
    st.subheader("최신 자소서 동향 + 기업 인재상/요구역량")
    c1, c2 = st.columns(2)

    # ----- 왼쪽 컬럼: 회사/직무 입력 -----
    with c1:
        t_company = st.text_input("회사명", key="trend_company")
        t_role = st.text_input("직무", key="trend_role")
        do_crawl = st.button("웹 리서치 실행")  # 웹 크롤링 버튼

    # ----- 오른쪽 컬럼: Tableau Public 링크 입력 -----
    with c2:
        tableau_link = st.text_input("(선택) Tableau Public 링크 임베드")

    # Tableau 링크가 있으면 iframe으로 표시
    if tableau_link:
        st.markdown("---")
        st.markdown("**Tableau Public**")
        st.components.v1.iframe(tableau_link, height=520)

    st.markdown("---")
    st.markdown("### 📊 로컬 데이터 인사이트")

    # 로컬 채용공고 데이터 시각화
    if PANDAS_OK and job_market is not None:
        if VIZ_OK and 'posted_date' in job_market.columns:
            try:
                jdf = job_market.copy()
                jdf['posted_date'] = pd.to_datetime(jdf['posted_date'], errors='coerce')
                ts = jdf.groupby(pd.Grouper(key='posted_date', freq='M'))['job_code'].nunique().reset_index()
                ts.columns = ['월', '공고수']
                st.altair_chart(
                    alt.Chart(ts).mark_line(point=True).encode(
                        x='월:T', 
                        y='공고수:Q'
                    ).properties(height=280, title='월별 채용공고 추이'),
                    use_container_width=True,
                )
            except Exception:
                pass

    # 상위 기술 수요 차트
    if PANDAS_OK and skills is not None and VIZ_OK:
        top_month = skills['month'].max() if 'month' in skills.columns else None
        kdf = skills.copy()
        if top_month:
            kdf = kdf[kdf['month']==top_month]
        if 'job_count' in kdf.columns:
            kdf = kdf.groupby('skill')['job_count'].sum().sort_values(ascending=False).head(15).reset_index()
            st.altair_chart(
                alt.Chart(kdf).mark_bar().encode(
                    x='job_count', 
                    y=alt.Y('skill', sort='-x')
                ).properties(height=360, title=f"{top_month or ''} 상위 기술 수요"),
                use_container_width=True,
            )

    # 웹 크롤링 실행 시 회사 인재상/요구역량 수집
    if do_crawl and t_company:
        if not HTTP_OK:
            st.warning("requests/bs4 미설치로 웹 리서치를 생략합니다. 'pip install requests beautifulsoup4' 설치 후 재시도")
        else:
            with st.spinner("회사 인재상/요구역량 수집 중…"):
                info = company_persona_and_requirements(t_company, t_role)  # 크롤링 결과
            # 인재상 표시
            if info.get("인재상"):
                st.markdown("### 🏢 인재상 요약")
                st.write(info["인재상"])
            # 요구역량 표시
            if info.get("요구역량"):
                st.markdown("### ✅ 요구역량(트렌드 기반 제안)")
                st.write(info["요구역량"])
            # 출처 표시 (최대 8개)
            if info.get("출처"):
                st.markdown("#### 🔗 참고 링크")
                for s in info["출처"][:8]:
                    st.markdown(f"- [{s.get('title','(제목없음)')}]({s.get('url','')}) — {s.get('snippet','')}")
