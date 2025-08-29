# =========================================================
# AI 자기소개서 코칭 - 통합 완성본 (v11 UI + v12 기능)
# =========================================================
# 설치: pip install streamlit python-docx reportlab langchain langchain-openai python-dotenv pandas numpy plotly
# 실행: streamlit run integrated_app.py
# =========================================================

import os, io, datetime, json, re, textwrap
from typing import Optional, List, Dict, Tuple
import streamlit as st

# ===== 문서 생성 라이브러리 (선택) =====
try:
    from docx import Document
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    DOC_LIBS_AVAILABLE = True
except:
    DOC_LIBS_AVAILABLE = False

# ===== LangChain (선택) =====
try:
    from langchain_openai import ChatOpenAI
    from langchain.prompts import ChatPromptTemplate
    from langchain.chains import LLMChain
    LANGCHAIN_AVAILABLE = True
except:
    LANGCHAIN_AVAILABLE = False

# ===== 데이터 분석 라이브러리 (선택) =====
try:
    import pandas as pd
    import numpy as np
    PANDAS_AVAILABLE = True
except:
    PANDAS_AVAILABLE = False

try:
    import plotly.express as px
    import plotly.graph_objects as go
    PLOTLY_AVAILABLE = True
except:
    PLOTLY_AVAILABLE = False

# ================= 세션 초기화 =================
if "messages" not in st.session_state:
    st.session_state.messages = []
    st.session_state.messages.append({
        "role": "ai",
        "content": "안녕하세요! AI 자기소개서 코치입니다. 무엇을 도와드릴까요?",
        "time": datetime.datetime.now().strftime("%H:%M")
    })

if "current_tab" not in st.session_state:
    st.session_state.current_tab = "대화"

if "api_key" not in st.session_state:
    st.session_state.api_key = os.getenv("OPENAI_API_KEY", "")

if "saved_files" not in st.session_state:
    st.session_state.saved_files = []

if "basic_settings" not in st.session_state:
    st.session_state.basic_settings = {
        "model": "GPT-4 (무료)",
        "tone": "전문적",
        "length": 800,
    }

if "advanced_settings" not in st.session_state:
    st.session_state.advanced_settings = {
        "creativity": 0.5,
        "polish": 0.5,
        "auto_save": True,
        "smart_edit": True,
        "export_format": "PDF 문서",
        "enable_scoring": True,
        "enable_trends": False,
    }

if "show_saved" not in st.session_state:
    st.session_state.show_saved = False

if "evaluation_history" not in st.session_state:
    st.session_state.evaluation_history = []

if "last_evaluation" not in st.session_state:
    st.session_state.last_evaluation = None

# ================= v12 스코어링 시스템 통합 =================
INDUSTRY_KEYWORDS = {
    "IT/개발": ["개발", "프로그래밍", "코드", "알고리즘", "데이터", "시스템", "API", "프레임워크"],
    "마케팅": ["캠페인", "브랜딩", "고객", "타겟", "채널", "ROI", "CPC", "CTR"],
    "영업": ["매출", "목표", "달성", "고객관계", "협상", "제안", "계약"],
    "기획": ["전략", "기획", "분석", "프로젝트", "로드맵", "KPI", "보고서"],
}

ACTION_WORDS = ["개선", "최적화", "설계", "구현", "분석", "자동화", "협업", "리팩터", "검증",
                "성과", "증가", "감소", "달성", "기여", "해결", "리더", "조율"]

STAR_TOKENS = ["상황", "과제", "행동", "결과", "Situation", "Task", "Action", "Result"]

FILLERS = ["최대한", "정말", "매우", "다양한", "많은", "열정", "성실", "노력"]

NUM_RE = re.compile(r"(?<!\w)(?:[0-9]+(?:\.[0-9]+)?%?|[일이삼사오육칠팔구십백천만]+%?)(?!\w)")

def tokenize_kr(text: str) -> List[str]:
    return re.findall(r"[\w가-힣%]+", text.lower())

def compute_resume_scores(text: str, role: str = "", company: str = "") -> Dict:
    tokens = tokenize_kr(text)
    n_words = len(tokens)
    n_chars = len(text)
    n_sentences = len(re.split(r'[.!?。]', text))
    
    # 1. 숫자(성과) 밀도
    nums = NUM_RE.findall(text)
    metric_density = min(1.0, len(nums) / max(1, n_words) * 10)
    
    # 2. 행동동사/액션
    action_hits = sum(1 for w in ACTION_WORDS if any(w in t for t in tokens))
    action_score = min(1.0, action_hits / 6)
    
    # 3. STAR 단서
    star_hits = sum(1 for w in STAR_TOKENS if any(w.lower() in t for t in tokens))
    star_score = min(1.0, star_hits / 4)
    
    # 4. 군더더기(감점)
    filler_hits = sum(tokens.count(f.lower()) for f in FILLERS)
    filler_penalty = min(0.3, filler_hits / max(1, n_words) * 5)
    
    # 5. 길이 적정성
    if 600 <= n_chars <= 1200:
        length_score = 1.0
    elif 400 <= n_chars < 600 or 1200 < n_chars <= 1500:
        length_score = 0.7
    else:
        length_score = 0.4
    
    # 6. 문장 다양성
    sentence_lengths = [len(s.strip()) for s in re.split(r'[.!?。]', text) if s.strip()]
    if len(sentence_lengths) > 1 and PANDAS_AVAILABLE:
        std_dev = np.std(sentence_lengths)
        variety_score = min(1.0, std_dev / 30)
    else:
        variety_score = 0.3
    
    # 가중합
    weights = {
        'metrics': 0.25,
        'action': 0.20,
        'star': 0.20,
        'length': 0.15,
        'variety': 0.10,
        'filler': -0.10,
    }
    
    total = (
        metric_density * weights['metrics'] +
        action_score * weights['action'] +
        star_score * weights['star'] +
        length_score * weights['length'] +
        variety_score * weights['variety'] +
        filler_penalty * weights['filler']
    )
    total = max(0.0, min(1.0, total))
    
    # 등급 부여
    if total >= 0.9:
        grade = "A+"
    elif total >= 0.8:
        grade = "A"
    elif total >= 0.7:
        grade = "B+"
    elif total >= 0.6:
        grade = "B"
    elif total >= 0.5:
        grade = "C+"
    else:
        grade = "C"
    
    return {
        '총점': round(total * 100, 1),
        '등급': grade,
        '성과밀도': round(metric_density * 100, 1),
        '행동성': round(action_score * 100, 1),
        'STAR구조': round(star_score * 100, 1),
        '길이적정': round(length_score * 100, 1),
        '문장다양성': round(variety_score * 100, 1),
        '군더더기': round(filler_penalty * 100, 1),
        '문장수': n_sentences,
        '단어수': n_words,
        '글자수': n_chars,
    }

# ================= 가이드라인 응답 =================
def get_guideline() -> str:
    return """📝 **AI 자기소개서 입력 가이드라인**

**1. 구체적으로 질문하기**
✅ "마케팅 직무 신입 자기소개서 도입부 작성해줘"
❌ "자소서 써줘"

**2. 배경 정보 제공하기**
• 지원 회사와 직무
• 본인의 주요 경험
• 강조하고 싶은 역량

**3. 효과적인 질문 예시**
• "고객 서비스 경험을 영업직무에 연결하는 방법"
• "프로젝트 경험을 STAR 기법으로 정리해줘"
• "IT 기업 지원동기 작성 도와줘"

**4. 첨삭 요청 방법**
• 작성한 문장을 복사 후 "이 내용 첨삭해줘"
• 파일 업로드 후 "구체성 높여줘"
• "이 문장 더 임팩트 있게 수정해줘"

**5. 단계별 접근**
1️⃣ 전체 구조 잡기
2️⃣ 각 문단 작성
3️⃣ 표현 다듬기
4️⃣ 최종 검토

💡 **Tip**: 한 번에 모든 걸 해결하려 하지 말고, 단계별로 질문하세요!"""

# ================= AI 응답 생성 (스코어링 기능 통합) =================
def get_ai_response(user_input: str, uploaded_file=None) -> str:
    guideline_keywords = ["가이드", "가이드라인", "도움말", "사용법", "어떻게"]
    if any(keyword in user_input for keyword in guideline_keywords):
        return get_guideline()
    
    # 평가 요청 감지
    eval_keywords = ["평가", "점수", "채점", "분석해", "평가해"]
    if any(keyword in user_input for keyword in eval_keywords) and st.session_state.advanced_settings.get("enable_scoring", False):
        if uploaded_file:
            try:
                if uploaded_file.name.endswith('.txt'):
                    content = uploaded_file.read().decode('utf-8')
                elif uploaded_file.name.endswith('.docx') and DOC_LIBS_AVAILABLE:
                    doc = Document(uploaded_file)
                    content = '\n'.join([p.text for p in doc.paragraphs])
                else:
                    content = "파일을 읽을 수 없습니다."
                
                scores = compute_resume_scores(content)
                st.session_state.last_evaluation = scores
                
                response = f"""📊 **자기소개서 평가 결과**

**총점: {scores['총점']}점 (등급: {scores['등급']})**

📈 **세부 점수**
• 성과 지표 밀도: {scores['성과밀도']}%
• 행동 동사 사용: {scores['행동성']}%
• STAR 구조: {scores['STAR구조']}%
• 길이 적정성: {scores['길이적정']}%
• 문장 다양성: {scores['문장다양성']}%
• 군더더기 (감점): -{scores['군더더기']}%

📝 **기본 정보**
• 글자 수: {scores['글자수']}자
• 단어 수: {scores['단어수']}개
• 문장 수: {scores['문장수']}개

💡 **개선 포인트**
"""
                if scores['성과밀도'] < 50:
                    response += "\n• 구체적인 숫자와 성과를 더 추가하세요"
                if scores['행동성'] < 50:
                    response += "\n• 행동 동사를 더 활용하세요 (개선, 구현, 달성 등)"
                if scores['STAR구조'] < 50:
                    response += "\n• STAR 구조로 경험을 재구성하세요"
                if scores['길이적정'] < 70:
                    response += "\n• 적정 길이(600-1200자)로 조정하세요"
                
                return response
            except Exception as e:
                return f"파일 처리 중 오류: {e}"
        else:
            return "평가할 자기소개서 파일을 업로드해주세요."

    if not st.session_state.api_key or not LANGCHAIN_AVAILABLE:
        templates = {
            "default": """자기소개서 작성을 도와드리겠습니다!

구체적으로 알려주시면 더 정확한 도움을 드릴 수 있어요:
• 어떤 직무에 지원하시나요?
• 어떤 부분이 어려우신가요?
• 특별히 강조하고 싶은 경험이 있나요?""",
            "첨삭": """자기소개서 첨삭 포인트를 알려드릴게요:

✅ 구체적인 숫자와 성과 포함
✅ 직무와 연관된 경험 강조
✅ 문장은 간결하고 명확하게
✅ 진정성 있는 지원동기

파일을 업로드하거나 내용을 보내주시면 더 자세히 봐드릴게요!""",
            "시작": """자기소개서 작성을 시작해볼까요?

**Step 1. 기본 정보**
• 지원 회사:
• 지원 직무:
• 경력 구분: (신입/경력)

이 정보를 알려주시면 맞춤형으로 도와드릴게요!""",
            "예시": """다음은 간단한 자기소개서 예시입니다:

"문제 해결 능력을 바탕으로 한 프로젝트 경험을 통해 팀에 기여했던 사례가 있습니다."

이와 같은 방식으로 경험을 구체적으로 설명해보세요!""",
        }
        if "첨삭" in user_input or "수정" in user_input:
            return templates["첨삭"]
        elif "시작" in user_input or "처음" in user_input:
            return templates["시작"]
        elif "예시" in user_input:
            return templates["예시"]
        else:
            return templates["default"]

    try:
        model_map = {
            "GPT-4 (무료)": "gpt-4o-mini",
            "GPT-4": "gpt-4o",
            "GPT-3.5": "gpt-3.5-turbo",
        }
        selected_model = st.session_state.basic_settings.get("model", "GPT-4 (무료)")
        model_name = model_map.get(selected_model, "gpt-4o-mini")
        llm = ChatOpenAI(
            api_key=st.session_state.api_key,
            model=model_name,
            temperature=st.session_state.advanced_settings["creativity"],
        )

        system_prompt = f"""당신은 전문 자기소개서 작성 코치입니다.
        톤: {st.session_state.basic_settings['tone']}
        최대 길이: {st.session_state.basic_settings['length']}자

        - 구체적이고 실용적인 조언
        - 예시를 들어 설명
        - 친근하면서도 전문적인 톤
        - 이모지는 최소한으로 사용"""

        if uploaded_file:
            try:
                if uploaded_file.name.endswith('.txt'):
                    content = uploaded_file.read().decode('utf-8')
                elif uploaded_file.name.endswith('.docx') and DOC_LIBS_AVAILABLE:
                    doc = Document(uploaded_file)
                    content = '\n'.join([p.text for p in doc.paragraphs])
                else:
                    content = "파일을 읽을 수 없습니다."
                user_input = f"다음 자기소개서를 검토하고 개선점을 제안해주세요:\n\n{content}\n\n{user_input}"
            except Exception as e:
                return f"파일 처리 중 오류: {e}"

        prompt = ChatPromptTemplate.from_messages([
            ("system", system_prompt),
            ("human", "{input}")
        ])
        chain = LLMChain(llm=llm, prompt=prompt)
        response = chain.invoke({"input": user_input})
        return response.get("text", str(response))
    except Exception as e:
        return f"오류가 발생했습니다. 다시 시도해주세요.\n{str(e)}"

# ================= 대화 저장 (평가 결과 포함) =================
def save_conversation():
    content = ""
    for msg in st.session_state.messages:
        role = "👤 사용자" if msg["role"] == "user" else "🤖 AI 코치"
        content += f"[{msg.get('time', '')}] {role}\n{msg['content']}\n\n"
    
    # 마지막 평가 결과 추가
    if st.session_state.last_evaluation:
        content += "\n\n--- 마지막 평가 결과 ---\n"
        content += json.dumps(st.session_state.last_evaluation, ensure_ascii=False, indent=2)

    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"자소서대화_{timestamp}"
    export = st.session_state.advanced_settings.get("export_format", "텍스트 파일")

    if export == "PDF 문서" and DOC_LIBS_AVAILABLE:
        bio = io.BytesIO()
        doc = SimpleDocTemplate(bio, pagesize=letter)
        styles = getSampleStyleSheet()
        story = [Paragraph(p, styles["Normal"]) for p in content.split('\n')]
        doc.build(story)
        file_data = bio.getvalue()
        mime = "application/pdf"
        ext = "pdf"
    elif export == "Word 문서" and DOC_LIBS_AVAILABLE:
        doc = Document()
        doc.add_heading('AI 자기소개서 코칭 대화', 0)
        for para in content.split('\n'):
            doc.add_paragraph(para)
        bio = io.BytesIO()
        doc.save(bio)
        file_data = bio.getvalue()
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext = "docx"
    elif export == "HTML 문서":
        file_data = f"<html><body><pre>{content}</pre></body></html>"
        mime = "text/html"
        ext = "html"
    else:
        file_data = content
        mime = "text/plain"
        ext = "txt"

    st.session_state.saved_files.append({
        "name": f"{filename}.{ext}",
        "date": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
        "size": len(file_data),
        "data": file_data,
        "mime": mime
    })

    return f"{filename}.{ext}"

# ================= 페이지 설정 및 기본 스타일 =================
st.set_page_config(
    page_title="AI 자기소개서 코칭",
    page_icon="💬",
    layout="wide",
    initial_sidebar_state="collapsed"
)

MAIN_COLOR = "#22C55E"       # 메인 초록색
SUB_COLOR = "#DCFCE7"        # 사용자 말풍선 배경
BOT_COLOR = "#F3F4F6"        # 챗봇 말풍선 배경
BG_COLOR = "#F5FBFB"         # 전체 배경색

st.markdown(
    f"""
    <style>
        body {{
            background-color: {BG_COLOR};
        }}

        .chat-header-title {{
            color: white;
            font-weight: 600;
        }}
        .bottom-nav {{
            position: fixed;
            left: 0;
            right: 0;
            bottom: 0;
            background: white;
            border-top: 1px solid #e0e0e0;
            padding: 4px 8px;
        }}
        .bottom-nav button {{
            width: 100%;
            background: transparent;
            border: none;
            color: {MAIN_COLOR};
            font-size: 14px;
        }}
        .bottom-nav .active {{
            color: white;
            background: {MAIN_COLOR};
            border-radius: 12px;
        }}
        .nav-icon {{
            font-size: 20px;
            display: block;
        }}
        .onboard-wrapper {{
            text-align: center;
            padding: 60px 20px;
        }}
        .onboard-circle {{
            width: 120px;
            height: 120px;
            border-radius: 60px;
            background: {SUB_COLOR};
            margin: 0 auto 24px auto;
            display:flex;
            align-items:center;
            justify-content:center;
            font-size:32px;
        }}
        .score-card {{
            background: white;
            border-radius: 12px;
            padding: 15px;
            margin: 10px 0;
            border: 1px solid #e0e0e0;
        }}
        .stMainBlockContainer {{
            // padding: 0;
        }}
        .stVerticalBlock {{
            // gap: 0;
        }}
        .stAppHeader,.stDecoration {{
            display: none;
        }}
        .header {{
            position: fixed;
            top: 0;
            left: 0;
            width: 100%;
            z-index: 10;
        }}
        .stFileUploader button {{
            display: none;
        }}
    </style>
    """,
    unsafe_allow_html=True,
)

# ================= UI 렌더링 함수 =================
def render_header(title: str) -> None:
    st.markdown(
        f"<div class='header' style='background:{MAIN_COLOR};padding:12px; text-align:center; color:white; font-weight:600'>{title}</div>",
        unsafe_allow_html=True,
    )

def render_bottom_nav() -> None:
    cols = st.columns(4)
    tabs = ["대화", "설정", "세부 설정", "계정"]
    icons = ["💬", "⚙️", "🛠️", "👤"]
    
    for col, tab, icon in zip(cols, tabs, icons):
        label = f"{icon} {tab}"
        if col.button(label, key=f"nav_{tab}", use_container_width=True):
            st.session_state.current_tab = tab
            st.rerun()

def render_onboarding():
    render_header("AI 자기소개서")
    st.markdown(
        "<div class='onboard-wrapper'>"\
        "<div class='onboard-circle'>✍️</div>"\
        "<h3>AI 자기소개서</h3>"\
        "<p>AI와 대화하면서 나만의 탄탄한 자기소개서를 완성하세요.</p>"\
        "<ol style='text-align:left; display:inline-block;'>"\
        "<li>AI와 대화를 통해 작성의 방향을 잡아</li>"\
        "<li>궁금한 질문은 언제든지 톡! 작성 톤을 설정하고</li>"\
        "<li>완벽하게 마무리된 자기소개서를 완성</li>"\
        "</ol>"\
        "</div>",
        unsafe_allow_html=True,
    )
    if st.button("시작하기", use_container_width=True):
        st.session_state.started = True
        st.session_state.current_tab = "대화"
        st.rerun()

def render_chat_tab():
    render_header("AI 대화")
    
    # 마지막 평가 결과 표시 (있는 경우)
    if st.session_state.last_evaluation and st.session_state.advanced_settings.get("enable_scoring", False):
        st.markdown(
            f"""<div class='score-card'>
            <b>📊 최근 평가</b> | 총점: {st.session_state.last_evaluation['총점']}점 | 등급: {st.session_state.last_evaluation['등급']}
            </div>""",
            unsafe_allow_html=True
        )
    
    for msg in st.session_state.messages:
        if msg["role"] == "user":
            st.markdown(
                f"<div style='text-align:right; background:{SUB_COLOR}; padding:10px; border-radius:18px; margin:4px 0'>{msg['content']}</div>",
                unsafe_allow_html=True,
            )
        else:
            content_html = msg["content"].replace("\n", "<br>")
            st.markdown(
                f"<div style='text-align:left; background:{BOT_COLOR}; padding:10px; border-radius:18px; margin:4px 0'>{content_html}</div>",
                unsafe_allow_html=True,
            )

    st.write("---")
    uploaded_file = st.file_uploader("📎 파일 첨부 (txt, docx)", type=["txt", "docx"])

    # --- 상태 초기화 ---
    st.session_state.setdefault("user_input", "")
    st.session_state.setdefault("_submit", False)
    st.session_state.setdefault("pending_input", None)

    # 콜백: 엔터/버튼 → 제출 의도 표시 + 입력 비우기(여기서만 비움)
    def submit_message():
        v = st.session_state.user_input.strip()
        if v:
            st.session_state.pending_input = v   # 본문에서 사용할 버퍼
            st.session_state.user_input = ""     # 위젯 값은 콜백에서만 리셋
            st.session_state._submit = True

    col1, col2, col3, col4 = st.columns([5, 1, 1, 1])
    with col1:
        st.text_input(
            "메시지",
            key="user_input",
            placeholder="메시지를 입력하세요... (평가/점수 요청 가능)",
            label_visibility="collapsed",
            on_change=submit_message,   # 엔터로 제출
        )
    with col2:
        st.button("전송", on_click=submit_message)  # 버튼 제출
    with col3:
        save = st.button("저장하기")
    with col4:
        if st.button("📂"):
            st.session_state.show_saved = not st.session_state.get("show_saved", False)

    # 제출 처리: 콜백이 남겨둔 pending_input을 사용
    if st.session_state._submit and st.session_state.pending_input:
        user_input = st.session_state.pending_input
        st.session_state._submit = False
        st.session_state.pending_input = None

        st.session_state.messages.append({
            "role": "user",
            "content": user_input,
            "time": datetime.datetime.now().strftime("%H:%M"),
        })
        with st.spinner("답변 생성 중..."):
            response = get_ai_response(user_input, uploaded_file)
        st.session_state.messages.append({
            "role": "ai",
            "content": response,
            "time": datetime.datetime.now().strftime("%H:%M"),
        })
        st.rerun()

    if save:
        filename = save_conversation()
        st.success(f
