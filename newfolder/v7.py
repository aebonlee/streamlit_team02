# app.py
# =========================================================
# Requirements (install first):
#   pip install -U pip
#   pip install streamlit python-docx reportlab langchain langchain-openai langchain-google-genai python-dotenv googletrans
#
# Optional:
#   - (PDF 한글 폰트) 프로젝트 폴더에 NanumGothic.ttf를 넣으면 PDF 한글이 깨지지 않아요.
#   - OPENAI_API_KEY / GEMINI_API_KEY는 .env에 넣거나, 화면의 설정 탭에서 직접 입력하세요.
# =========================================================

import os, io, json, time, textwrap, re, datetime, urllib.parse, base64
from typing import Optional, Tuple, List, Dict

import streamlit as st

# ===== 문서 생성을 위한 라이브러리 =====
try:
    from docx import Document
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    DOC_LIBS_AVAILABLE = True
except Exception:
    DOC_LIBS_AVAILABLE = False

# ===== LangChain imports (조건부) =====
try:
    from langchain_openai import ChatOpenAI
    from langchain_google_genai import ChatGoogleGenerativeAI
    from langchain.prompts import ChatPromptTemplate, MessagesPlaceholder
    from langchain.memory import ConversationBufferMemory
    from langchain.chains import LLMChain
    LANGCHAIN_AVAILABLE = True
except Exception:
    LANGCHAIN_AVAILABLE = False

# ===== 번역 라이브러리 (옵션) =====
try:
    from googletrans import Translator
    TRANSLATOR_AVAILABLE = True
except Exception:
    TRANSLATOR_AVAILABLE = False

# ================= 기본 설정 =================
st.set_page_config(
    page_title="AI 자기소개서 코치",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ================= 스타일 =================
st.markdown("""
<style>
.main .block-container { max-width: 900px; padding: 2rem 1rem; }
.header-container {
  background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
  border-radius: 20px; padding: 2rem; margin-bottom: 2rem;
  text-align: center; color: white; box-shadow: 0 10px 30px rgba(0,0,0,0.1);
}
.header-title { font-size: 2.5rem; font-weight: 800; margin-bottom: 0.5rem; }
.header-subtitle { font-size: 1.1rem; opacity: 0.9; font-weight: 300; }
.chat-container {
  background: white; border-radius: 20px; padding: 1.5rem; margin-bottom: 1rem;
  box-shadow: 0 5px 20px rgba(0,0,0,0.05); border: 1px solid #f0f0f0;
  height: 500px; overflow-y: auto;
}
.message-bubble { margin: 1rem 0; display: flex; align-items: flex-start; }
.message-bubble.user { justify-content: flex-end; }
.message-bubble.bot { justify-content: flex-start; }
.message-content {
  max-width: 75%; padding: 1rem 1.5rem; border-radius: 20px; font-size: 0.95rem;
  line-height: 1.6; word-wrap: break-word;
}
.message-bubble.user .message-content {
  background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
  color: white; border-bottom-right-radius: 5px;
}
.message-bubble.bot .message-content {
  background: #f8fafc; color: #1a202c; border: 1px solid #e2e8f0; border-bottom-left-radius: 5px;
}
.stTabs [data-baseweb="tab-list"] {
  gap: 0.5rem; background: #f8fafc; padding: 0.5rem; border-radius: 15px;
}
.stTabs [data-baseweb="tab"] {
  height: 3rem; padding: 0 1.5rem; background: white; border-radius: 10px; border: none; font-weight: 500;
}
.stTabs [aria-selected="true"] {
  background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white;
}
.feature-card, .file-item {
  background: white; border-radius: 15px; padding: 1.0rem; margin: 0.5rem 0;
  box-shadow: 0 5px 15px rgba(0,0,0,0.05); border: 1px solid #e2e8f0;
}
.guideline-card {
  background: linear-gradient(135deg, #e0f2f1 0%, #f3e5f5 100%); border-radius: 15px; padding: 1.5rem;
  margin: 1rem 0; border-left: 4px solid #667eea;
}
.guideline-item {
  background: white; border-radius: 10px; padding: 1rem; margin: 0.5rem 0;
  box-shadow: 0 2px 8px rgba(0,0,0,0.05);
}
.small { font-size: 0.85rem; color: #4b5563; }
</style>
""", unsafe_allow_html=True)

# ================= 상태 초기화 =================
def _now_hhmm():
    return datetime.datetime.now().strftime("%H:%M")

def _timestamp():
    return datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

if "initialized" not in st.session_state:
    st.session_state.initialized = True
    st.session_state.msgs = []
    st.session_state.settings = {
        "provider": "openai",
        "model": "gpt-4o-mini",        # gemini 사용 시 아래에서 강제로 "gemini-1.5-pro"로 치환
        "tone": "정중하고 간결한",
        "length": 800,
        "temperature": 0.7,
        "openai_key": os.getenv("OPENAI_API_KEY", ""),
        "gemini_key": os.getenv("GEMINI_API_KEY", ""),
        "save_dir": "./AI_CoverLetter_Storage",
        "font_family": "NanumGothic",
        "enable_translation": False,
        "use_free_model": True
    }
    os.makedirs(st.session_state.settings["save_dir"], exist_ok=True)
    st.session_state.msgs.append({
        "role": "bot",
        "content": "안녕하세요! AI 자기소개서 코치입니다. 🎯\n\n어떤 도움이 필요하신가요?",
        "timestamp": _now_hhmm()
    })

if "saved_files" not in st.session_state:
    st.session_state.saved_files = []

if LANGCHAIN_AVAILABLE and "lc_memory" not in st.session_state:
    st.session_state.lc_memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)

# ================= 유틸 함수 =================
def translate_to_english(text: str) -> str:
    if not TRANSLATOR_AVAILABLE:
        return "번역 기능을 사용하려면 googletrans 라이브러리를 설치해주세요."
    try:
        translator = Translator()
        return translator.translate(text, src='ko', dest='en').text
    except Exception as e:
        return f"번역 중 오류가 발생했습니다: {e}"

def _ensure_korean_font(font_path: str, font_name: str = "NanumGothic"):
    """ReportLab에 한글 폰트를 등록 (없으면 기본 폰트 사용)"""
    if not DOC_LIBS_AVAILABLE:
        return None
    if not os.path.exists(font_path):
        return None
    try:
        pdfmetrics.registerFont(TTFont(font_name, font_path))
        return font_name
    except Exception:
        return None

def _read_uploaded_text(uploaded_file) -> str:
    """txt/docx 업로드 파일을 안전하게 텍스트로 파싱"""
    name = uploaded_file.name.lower()
    if name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8", errors="ignore")
    elif name.endswith(".docx"):
        if not DOC_LIBS_AVAILABLE:
            raise RuntimeError("DOCX를 처리하려면 python-docx가 필요합니다.")
        # python-docx는 파일 객체도 직접 열 수 있습니다.
        _doc = Document(uploaded_file)
        return "\n".join(p.text for p in _doc.paragraphs)
    else:
        raise RuntimeError("지원하지 않는 파일 형식입니다. txt 또는 docx만 업로드하세요.")

def get_free_ai_response(user_message: str) -> str:
    response_templates = {
        "마케팅": """📊 **마케팅 직무 자기소개서 작성 가이드**

**1. 핵심 역량 강조**
- 데이터 분석 및 인사이트 도출 능력
- 창의적 캠페인 기획 경험
- 디지털 마케팅 도구 활용 능력

**2. 구체적 성과 제시**
- "매출 20% 증가" 같은 정량적 결과
- "CTR 3% 향상" 등 구체 지표
- "신규 고객 1,000명 확보" 등 수치화

**3. 경험 서술 방법**
- STAR 기법(상황-과제-행동-결과)
- 문제 해결 과정과 결과 중심
- 팀워크/리더십 포함""",
        "개발": """💻 **개발 직무 자기소개서 작성 가이드**

**1. 기술 스택 명시**
- 언어/프레임워크/라이브러리
- DB/클라우드/CI-CD 경험

**2. 프로젝트 경험 상세화**
- 서비스 규모/성과
- 해결한 기술적 이슈와 접근
- 코드 품질/테스트/리팩토링 노력

**3. 성장 의지**
- 지속 학습/트렌드 관심
- 오픈소스/개인 프로젝트""",
        "영업": """🎯 **영업 직무 자기소개서 작성 가이드**

**1. 성과 강조**
- 목표 달성률/매출 기여
- 신규 고객 개발/리텐션

**2. 커뮤니케이션**
- 니즈 파악/솔루션 제안
- 프레젠테이션/협업 경험

**3. 시장 이해**
- 트렌드/경쟁사 분석
- 고객사 비즈니스 이해"""
    }
    user_lower = user_message.lower()
    if "마케팅" in user_lower:
        return response_templates["마케팅"]
    elif any(word in user_lower for word in ["개발", "프로그래밍", "코딩", "it"]):
        return response_templates["개발"]
    elif "영업" in user_lower:
        return response_templates["영업"]
    elif any(word in user_lower for word in ["첨삭", "피드백", "검토"]):
        return """✏️ **자기소개서 첨삭 포인트**

**1. 구조/논리**
- 도입-본론-결론
- 문단 간 연결/일관성

**2. 내용 구체성**
- 추상→사례, 수치화
- 차별화 포인트

**3. 문장 표현**
- 간결/자연스러운 어조
- 중복/군더더기 제거

📎 파일을 업로드하시면 더 구체적으로 도와드려요."""
    else:
        return """🎯 **자기소개서 작성을 도와드릴게요!**

**효과적인 질문 예시**
- "마케팅 직무 자기소개서 작성법"
- "개발자 자소서 핵심 포인트?"
- "영업 경험을 임팩트 있게 쓰는 법?"
- "제 자기소개서 첨삭해주세요"(파일 첨부)

**작성 원칙**
1) STAR 기법  2) 수치화  3) 차별화"""

def get_ai_response(user_message: str, uploaded_file=None) -> str:
    settings = st.session_state.settings

    # 무료 모드 또는 키 없음 → 템플릿 응답
    if settings["use_free_model"] or (not settings["openai_key"] and not settings["gemini_key"]):
        if uploaded_file is not None:
            try:
                file_content = _read_uploaded_text(uploaded_file)
                return f"""📋 **업로드된 자기소개서 첨삭(요약 미리보기)**

**원문 일부:**
{file_content[:200]}...

**첨삭 가이드:**
{get_free_ai_response("첨삭")}

💡 더 정교한 첨삭은 '설정'에서 API 키를 입력 후 사용하세요."""
            except Exception as e:
                return f"파일 읽기 오류: {e}"
        return get_free_ai_response(user_message)

    if not LANGCHAIN_AVAILABLE:
        return get_free_ai_response(user_message)

    # LLM 선택
    try:
        if settings["provider"] == "openai" and settings["openai_key"]:
            llm = ChatOpenAI(
                api_key=settings["openai_key"],
                model=settings["model"],
                temperature=settings["temperature"]
            )
        elif settings["provider"] == "gemini" and settings["gemini_key"]:
            # 최신 추천 모델명
            llm = ChatGoogleGenerativeAI(
                google_api_key=settings["gemini_key"],
                model="gemini-1.5-pro",
                temperature=settings["temperature"]
            )
        else:
            return get_free_ai_response(user_message)

        # 프롬프트 구성
        system_prompt = f"""
당신은 자기소개서 전문 코치입니다.
- 톤: {settings["tone"]}
- 목표 길이: 약 {settings["length"]}자
- 구체적/실용적 조언
- STAR 기법 권장
- 정량적 성과/구체 사례 강조
        """.strip()

        if uploaded_file is not None:
            try:
                file_content = _read_uploaded_text(uploaded_file)
            except Exception as e:
                return f"파일 읽기 오류: {e}"
            prompt_text = f"""
다음 자기소개서를 전문가 관점에서 첨삭해주세요.

[자기소개서]
{file_content}

[사용자 질문]
{user_message}

다음 관점에서 상세 피드백:
1) 구조/논리  2) 구체성/차별화  3) 표현/어법  4) 개선 제안
""".strip()
        else:
            prompt_text = user_message

        prompt = ChatPromptTemplate.from_messages([
            ("system", system_prompt),
            MessagesPlaceholder(variable_name="chat_history"),
            ("human", "{input}")
        ])

        chain = LLMChain(llm=llm, prompt=prompt, memory=st.session_state.lc_memory)

        # invoke를 사용하면 버전 차이로 인한 run 디프리케이션 이슈를 피할 수 있어요
        result = chain.invoke({"input": prompt_text})
        response_text = result.get("text") if isinstance(result, dict) else str(result)

        if settings["enable_translation"] and uploaded_file is None:
            eng = translate_to_english(response_text)
            response_text += f"\n\n---\n**영문 버전:**\n{eng}"

        return response_text

    except Exception as e:
        return f"오류가 발생했습니다: {e}\n\n{get_free_ai_response(user_message)}"

# ================= 문서 생성/저장 =================
def create_txt(content: str, filename: str) -> Optional[str]:
    try:
        filepath = os.path.join(st.session_state.settings["save_dir"], f"{filename}.txt")
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(f"자기소개서\n{'='*20}\n\n")
            f.write(content)
        return filepath
    except Exception as e:
        st.error(f"TXT 생성 중 오류: {e}")
        return None

def create_docx(content: str, filename: str) -> Optional[str]:
    if not DOC_LIBS_AVAILABLE:
        st.error("DOCX 생성에는 python-docx가 필요합니다. `pip install python-docx`")
        return None
    try:
        filepath = os.path.join(st.session_state.settings["save_dir"], f"{filename}.docx")
        doc = Document()
        title = doc.add_heading('자기소개서', 0)
        title.alignment = 1
        for line in content.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
            else:
                doc.add_paragraph("")
        doc.save(filepath)
        return filepath
    except Exception as e:
        st.error(f"DOCX 생성 중 오류: {e}")
        return None

def create_pdf(content: str, filename: str) -> Optional[str]:
    if not DOC_LIBS_AVAILABLE:
        st.error("PDF 생성에는 reportlab이 필요합니다. `pip install reportlab`")
        return None
    try:
        filepath = os.path.join(st.session_state.settings["save_dir"], f"{filename}.pdf")
        doc = SimpleDocTemplate(filepath, pagesize=letter)

        # 한글 폰트 적용 시도 (없으면 기본폰트)
        font_name = _ensure_korean_font("./NanumGothic.ttf") or "Helvetica"

        styles = getSampleStyleSheet()
        story = []

        # 제목 스타일
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1,
            fontName=font_name
        )
        normal_style = styles['Normal']
        normal_style.fontName = font_name

        story.append(Paragraph("자기소개서", title_style))
        story.append(Spacer(1, 12))

        for line in content.split('\n'):
            if line.strip():
                story.append(Paragraph(line, normal_style))
            else:
                story.append(Spacer(1, 6))

        doc.build(story)
        return filepath
    except Exception as e:
        st.error(f"PDF 생성 중 오류: {e}")
        return None

def _conversation_to_text() -> str:
    lines = []
    for msg in st.session_state.msgs:
        role = "👤 사용자" if msg["role"] == "user" else "🤖 AI 코치"
        lines.append(f"{role}: {msg['content']}\n")
        lines.append("---\n\n")
    return "".join(lines)

def save_conversation(file_format: str, filename: str) -> Optional[str]:
    text = _conversation_to_text()
    if file_format == "pdf":
        path = create_pdf(text, filename)
    elif file_format == "docx":
        path = create_docx(text, filename)
    else:
        path = create_txt(text, filename)

    if path and os.path.exists(path):
        info = {
            "name": os.path.basename(path),
            "path": path,
            "created_ts": os.path.getctime(path),
            "created": datetime.datetime.fromtimestamp(os.path.getctime(path)).strftime("%Y-%m-%d %H:%M:%S"),
            "size": os.path.getsize(path),
        }
        if info not in st.session_state.saved_files:
            st.session_state.saved_files.append(info)
        return path
    return None

def get_saved_files() -> List[Dict]:
    saved_files = []
    d = st.session_state.settings["save_dir"]
    if os.path.exists(d):
        for filename in os.listdir(d):
            path = os.path.join(d, filename)
            if os.path.isfile(path):
                ctime = os.path.getctime(path)
                saved_files.append({
                    "name": filename,
                    "path": path,
                    "created_ts": ctime,
                    "created": datetime.datetime.fromtimestamp(ctime).strftime("%Y-%m-%d %H:%M:%S"),
                    "size": os.path.getsize(path),
                })
    saved_files.sort(key=lambda x: x["created_ts"], reverse=True)
    return saved_files

# ================= UI 컴포넌트 =================
def render_header():
    st.markdown("""
    <div class="header-container">
        <div class="header-title">🎯 AI 자기소개서 코치</div>
        <div class="header-subtitle">전문 AI가 도와드리는 맞춤형 자기소개서 작성 서비스</div>
    </div>
    """, unsafe_allow_html=True)

def render_guidelines():
    st.markdown("""
    <div class="guideline-card">
        <h3>💡 효과적인 질문 가이드</h3>
        <div class="guideline-item"><strong>🎯 직무별</strong>
        <p>• "마케팅 직무 자소서 작성법"<br>• "개발자 자소서 핵심 포인트?"<br>• "영업 경험을 어떻게 어필?"</p></div>
        <div class="guideline-item"><strong>📝 상황별</strong>
        <p>• "신입 프로젝트 경험 작성"<br>• "경력 이직 사유 표현"<br>• "타분야 전환 포인트"</p></div>
        <div class="guideline-item"><strong>✏️ 기법</strong>
        <p>• "STAR 기법 구조화"<br>• "성과 수치화 방법"<br>• "적정 길이/톤"</p></div>
        <div class="guideline-item"><strong>🔍 첨삭</strong>
        <p>• "제 자소서 첨삭해주세요"(파일 첨부)</p></div>
    </div>
    """, unsafe_allow_html=True)

def render_chat_tab():
    render_header()

    with st.expander("💡 질문 가이드라인 보기", expanded=False):
        render_guidelines()

    st.markdown('<div class="chat-container">', unsafe_allow_html=True)
    for msg in st.session_state.msgs:
        role_class = "user" if msg["role"] == "user" else "bot"
        st.markdown(f"""
        <div class="message-bubble {role_class}">
            <div class="message-content">{msg["content"].replace(chr(10), "<br>")}</div>
        </div>
        """, unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

    allowed_types = ['txt'] + (['docx'] if DOC_LIBS_AVAILABLE else [])
    uploaded_file = st.file_uploader(
        "📎 자기소개서 파일 첨부 (첨삭용)",
        type=allowed_types,
        help="TXT 또는 DOCX 파일을 업로드하면 첨삭을 도와드립니다."
    )

    with st.form(key="chat_form", clear_on_submit=True):
        c1, c2 = st.columns([5, 1])
        with c1:
            user_input = st.text_input("메시지를 입력하세요...",
                placeholder="예: 마케팅 직무 자기소개서 작성법을 알려주세요",
                label_visibility="collapsed")
        with c2:
            submit = st.form_submit_button("전송", use_container_width=True, type="primary")

        if submit and user_input.strip():
            st.session_state.msgs.append({"role": "user", "content": user_input.strip(), "timestamp": _now_hhmm()})
            with st.spinner("AI가 답변을 생성중입니다..."):
                ai_response = get_ai_response(user_input.strip(), uploaded_file)
            st.session_state.msgs.append({"role": "bot", "content": ai_response, "timestamp": _now_hhmm()})
            st.rerun()

    # ===== 대화 저장/다운로드 =====
    st.markdown("### 💾 대화 저장")
    c1, c2, c3 = st.columns([2,2,3])
    with c1:
        fmt = st.selectbox("파일 형식", ["txt", "docx", "pdf"], index=0)
    with c2:
        default_name = f"conversation_{_timestamp()}"
        filename = st.text_input("파일명", value=default_name)
    with c3:
        if st.button("저장하기", type="secondary"):
            path = save_conversation(fmt, filename)
            if path:
                st.success(f"저장됨: {path}")
                # 즉시 다운로드 버튼 제공
                with open(path, "rb") as f:
                    st.download_button("💾 파일 다운로드", f, file_name=os.path.basename(path))

    # 저장된 파일 리스트
    st.markdown("### 📂 저장된 파일")
    files = get_saved_files()
    if not files:
        st.caption("아직 저장된 파일이 없습니다.")
    else:
        for fobj in files[:10]:
            with st.container():
                st.markdown(
                    f"""<div class="file-item">
<b>{fobj["name"]}</b><br>
<span class="small">생성: {fobj["created"]} · 크기: {fobj["size"]} bytes</span>
</div>""",
                    unsafe_allow_html=True,
                )
                with open(fobj["path"], "rb") as fh:
                    st.download_button("다운로드", fh, file_name=fobj["name"], key=f"dl_{fobj['name']}")

def render_settings_tab():
    st.markdown('<div class="feature-card"><h2>⚙️ AI 모델 및 응답 설정</h2></div>', unsafe_allow_html=True)
    s = st.session_state.settings

    st.markdown("### 🆓 모델 사용 방식")
    use_free = st.checkbox("무료 모드 사용 (API 키 없이 기본 가이드 제공)", value=s["use_free_model"])
    s["use_free_model"] = use_free

    if not use_free:
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("### 🔑 API 키 설정")
            provider = st.selectbox("AI 제공업체", ["openai", "gemini"], index=0 if s["provider"] == "openai" else 1)
            s["provider"] = provider

            if provider == "openai":
                openai_key = st.text_input("OpenAI API Key", value=s["openai_key"], type="password", help="환경변수 OPENAI_API_KEY 사용 가능")
                s["openai_key"] = openai_key

                models = ["gpt-4o-mini", "gpt-4o", "gpt-3.5-turbo"]
                try:
                    idx = models.index(s["model"]) if s["model"] in models else 0
                except Exception:
                    idx = 0
                model = st.selectbox("OpenAI 모델", models, index=idx)
                s["model"] = model
            else:
                gemini_key = st.text_input("Google Gemini API Key", value=s["gemini_key"], type="password", help="환경변수 GEMINI_API_KEY 사용 가능")
                s["gemini_key"] = gemini_key
                s["model"] = "gemini-1.5-pro"

        with col2:
            st.markdown("### 🗣️ 응답 스타일")
            s["tone"] = st.selectbox("톤", ["정중하고 간결한", "자신감 있고 설득적인", "따뜻하고 공감적인"], index=0)
            s["length"] = st.slider("목표 길이(자)", min_value=200, max_value=2000, value=s["length"], step=50)
            s["temperature"] = st.slider("창의성(Temperature)", min_value=0.0, max_value=1.5, value=float(s["temperature"]), step=0.1)

        st.markdown("### 🌐 부가 기능")
        s["enable_translation"] = st.checkbox("응답을 영어로도 제공(번역)", value=s["enable_translation"])

    st.markdown("### 📁 저장 폴더")
    save_dir = st.text_input("저장 경로", value=s["save_dir"])
    if save_dir != s["save_dir"]:
        s["save_dir"] = save_dir
        os.makedirs(s["save_dir"], exist_ok=True)
        st.success(f"저장 폴더 적용: {s['save_dir']}")

    st.caption("※ PDF 한글 깨짐 시 프로젝트 폴더에 NanumGothic.ttf를 넣어 주세요.")

# ================= 메인 탭 =================
tab1, tab2 = st.tabs(["💬 채팅", "⚙️ 설정"])
with tab1:
    render_chat_tab()
with tab2:
    render_settings_tab()
