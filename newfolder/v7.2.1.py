# app.py
# =========================================================
# AI 자기소개서 코칭 - 카카오톡 스타일 UI
# =========================================================
# 설치: pip install streamlit python-docx reportlab langchain langchain-openai python-dotenv
# 실행: streamlit run appv72.py
# =========================================================

import os, io, datetime, json
from typing import Optional, List, Dict
import streamlit as st

# ===== 문서 생성 라이브러리 (선택) =====
try:
    from docx import Document
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    DOC_LIBS_AVAILABLE = True
except:
    DOC_LIBS_AVAILABLE = False

# ===== LangChain (선택) =====
try:
    from langchain_openai import ChatOpenAI
    from langchain.prompts import ChatPromptTemplate
    from langchain.chains import LLMChain
    LANGCHAIN_AVAILABLE = True
except:
    LANGCHAIN_AVAILABLE = False

# ================= 페이지 설정 =================
st.set_page_config(
    page_title="AI 자기소개서 코칭",
    page_icon="💬",
    layout="wide",
    initial_sidebar_state="collapsed"
)

# ================= 카카오톡 스타일 CSS =================
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+KR:wght@300;400;500;600;700&display=swap');
    
    /* 전체 배경 */
    .stApp {
        background: #b2c7d9;
    }
    
    /* 메인 컨테이너 */
    .main .block-container {
        padding: 0;
        max-width: 100%;
        margin: 0;
    }
    
    /* 상단 헤더 */
    .chat-header {
        background: rgba(0, 0, 0, 0.85);
        color: white;
        padding: 15px 20px;
        text-align: center;
        border-radius: 0 0 15px 15px;
        margin-bottom: 20px;
        backdrop-filter: blur(10px);
        border-bottom: 1px solid rgba(255, 255, 255, 0.1);
    }
    
    .header-title {
        font-size: 20px;
        font-weight: 600;
        color: white;
        margin: 0;
    }
    
    /* 채팅 영역 */
    .chat-container {
        background: white;
        border-radius: 15px;
        padding: 20px;
        margin-bottom: 20px;
        min-height: 400px;
        max-height: 600px;
        overflow-y: auto;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    
    /* 메시지 버블 */
    .msg-row {
        display: flex;
        margin-bottom: 15px;
        align-items: flex-end;
    }
    
    .msg-row.user {
        justify-content: flex-end;
    }
    
    .msg-row.ai {
        justify-content: flex-start;
    }
    
    .msg-bubble {
        max-width: 70%;
        padding: 12px 16px;
        border-radius: 18px;
        font-size: 14px;
        line-height: 1.5;
        word-break: break-word;
        box-shadow: 0 2px 4px rgba(0, 0, 0, 0.1);
        position: relative;
    }
    
    .msg-bubble.user {
        background: #ffeb33;
        color: #000;
        border-top-right-radius: 4px;
    }
    
    .msg-bubble.ai {
        background: #f8f9fa;
        color: #000;
        border-top-left-radius: 4px;
        border: 1px solid #e9ecef;
    }
    
    .msg-time {
        font-size: 11px;
        color: #888;
        margin: 0 8px;
        white-space: nowrap;
    }
    
    /* 입력창 영역 */
    .input-container {
        background: white;
        padding: 20px;
        border-radius: 15px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
        margin-bottom: 20px;
    }
    
    /* 빠른 답변 버튼 */
    .quick-reply {
        display: inline-block;
        padding: 8px 16px;
        margin: 4px;
        background: #f8f9fa;
        border: 1px solid #e9ecef;
        border-radius: 20px;
        font-size: 13px;
        cursor: pointer;
        transition: all 0.2s;
        text-decoration: none;
        color: #495057;
    }
    
    .quick-reply:hover {
        background: #ffeb33;
        border-color: #ffeb33;
        color: #000;
    }
    
    /* 설정 페이지 */
    .settings-container {
        background: white;
        border-radius: 15px;
        padding: 20px;
        margin-bottom: 20px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    
    .settings-section {
        background: #f8f9fa;
        border-radius: 12px;
        padding: 20px;
        margin-bottom: 15px;
        border: 1px solid #e9ecef;
    }
    
    .settings-title {
        font-size: 16px;
        font-weight: 600;
        margin-bottom: 15px;
        color: #333;
    }
    
    /* 저장소 페이지 */
    .storage-container {
        background: white;
        border-radius: 15px;
        padding: 20px;
        margin-bottom: 20px;
        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.1);
    }
    
    .file-item {
        background: #f8f9fa;
        border-radius: 10px;
        padding: 15px;
        margin-bottom: 10px;
        display: flex;
        justify-content: space-between;
        align-items: center;
        border: 1px solid #e9ecef;
    }
    
    .file-info {
        flex: 1;
    }
    
    .file-name {
        font-weight: 500;
        margin-bottom: 5px;
        color: #495057;
    }
    
    .file-date {
        font-size: 12px;
        color: #6c757d;
    }
    
    /* 버튼 스타일 */
    .stButton > button {
        background: #ffeb33;
        color: #000;
        border: none;
        border-radius: 20px;
        padding: 8px 20px;
        font-weight: 500;
        transition: all 0.2s;
    }
    
    .stButton > button:hover {
        background: #ffd900;
        transform: translateY(-1px);
    }
    
    /* 입력창 스타일 */
    .stTextInput > div > div > input {
        background: #f8f9fa;
        border: 1px solid #e9ecef;
        border-radius: 20px;
        padding: 12px 16px;
        font-size: 14px;
    }
    
    .stTextInput > div > div > input:focus {
        border-color: #ffeb33;
        box-shadow: 0 0 0 3px rgba(255, 235, 51, 0.1);
    }
    
    /* 파일 업로드 영역 */
    .stUploadedFile {
        background: #f8f9fa;
        border-radius: 10px;
        padding: 10px;
        margin-bottom: 10px;
        border: 1px solid #e9ecef;
    }
    
    /* selectbox 스타일 */
    .stSelectbox > div > div {
        background: #f8f9fa;
        border-radius: 10px;
    }
    
    /* 탭 스타일 */
    .stTabs [data-baseweb="tab-list"] {
        gap: 8px;
    }
    
    .stTabs [data-baseweb="tab"] {
        height: 50px;
        background: #f8f9fa;
        border-radius: 10px 10px 0 0;
        border: 1px solid #e9ecef;
        border-bottom: none;
        color: #6c757d;
    }
    
    .stTabs [aria-selected="true"] {
        background: #ffeb33;
        color: #000;
        border-color: #ffeb33;
    }
    
    /* 스크롤바 */
    ::-webkit-scrollbar {
        width: 8px;
    }
    
    ::-webkit-scrollbar-track {
        background: #f1f1f1;
        border-radius: 4px;
    }
    
    ::-webkit-scrollbar-thumb {
        background: #c1c1c1;
        border-radius: 4px;
    }
    
    ::-webkit-scrollbar-thumb:hover {
        background: #a8a8a8;
    }
</style>
""", unsafe_allow_html=True)

# ================= 세션 초기화 =================
if "messages" not in st.session_state:
    st.session_state.messages = []
    st.session_state.messages.append({
        "role": "ai",
        "content": "안녕하세요! AI 자기소개서 코치입니다. 무엇을 도와드릴까요?",
        "time": datetime.datetime.now().strftime("%H:%M")
    })

if "current_tab" not in st.session_state:
    st.session_state.current_tab = "대화"

if "api_key" not in st.session_state:
    st.session_state.api_key = os.getenv("OPENAI_API_KEY", "")

if "saved_files" not in st.session_state:
    st.session_state.saved_files = []

if "save_format" not in st.session_state:
    st.session_state.save_format = "txt"

if "model_settings" not in st.session_state:
    st.session_state.model_settings = {
        "temperature": 0.7,
        "max_length": 1000,
        "tone": "professional"
    }

# ================= 가이드라인 응답 =================
def get_guideline_response():
    return """📝 **AI 자기소개서 입력 가이드라인**

**1. 구체적으로 질문하기**
✅ "마케팅 직무 신입 자기소개서 도입부 작성해줘"
❌ "자소서 써줘"

**2. 배경 정보 제공하기**
• 지원 회사와 직무
• 본인의 주요 경험
• 강조하고 싶은 역량

**3. 효과적인 질문 예시**
• "고객 서비스 경험을 영업직무에 연결하는 방법"
• "프로젝트 경험을 STAR 기법으로 정리해줘"
• "IT 기업 지원동기 작성 도와줘"

**4. 첨삭 요청 방법**
• 작성한 내용 복사 후 "이 내용 첨삭해줘"
• 파일 업로드 후 "구체성 높여줘"
• "이 문장 더 임팩트 있게 수정해줘"

**5. 단계별 접근**
1️⃣ 전체 구조 잡기
2️⃣ 각 문단 작성
3️⃣ 표현 다듬기
4️⃣ 최종 검토

💡 **Tip**: 한 번에 모든 걸 해결하려 하지 말고, 단계별로 질문하세요!"""

# ================= AI 응답 생성 =================
def get_ai_response(user_input: str, uploaded_file=None) -> str:
    # 가이드라인 요청 체크
    guideline_keywords = ["가이드", "가이드라인", "도움말", "사용법", "어떻게"]
    if any(keyword in user_input for keyword in guideline_keywords):
        return get_guideline_response()
    
    # 템플릿 응답 (API 키 없을 때)
    if not st.session_state.api_key or not LANGCHAIN_AVAILABLE:
        templates = {
            "default": """자기소개서 작성을 도와드리겠습니다!

구체적으로 알려주시면 더 정확한 도움을 드릴 수 있어요:
• 어떤 직무에 지원하시나요?
• 어떤 부분이 어려우신가요?
• 특별히 강조하고 싶은 경험이 있나요?""",
            
            "첨삭": """자기소개서 첨삭 포인트를 알려드릴게요:

✅ 구체적인 숫자와 성과 포함
✅ 직무와 연관된 경험 강조
✅ 문장은 간결하고 명확하게
✅ 진정성 있는 지원동기

파일을 업로드하거나 내용을 보내주시면 더 자세히 봐드릴게요!""",
            
            "시작": """자기소개서 작성을 시작해볼까요?

**Step 1. 기본 정보**
• 지원 회사:
• 지원 직무:
• 경력 구분: (신입/경력)

이 정보를 알려주시면 맞춤형으로 도와드릴게요!"""
        }
        
        if "첨삭" in user_input or "수정" in user_input:
            return templates["첨삭"]
        elif "시작" in user_input or "처음" in user_input:
            return templates["시작"]
        else:
            return templates["default"]
    
    # LangChain을 이용한 AI 응답
    try:
        llm = ChatOpenAI(
            api_key=st.session_state.api_key,
            model="gpt-4o-mini",
            temperature=st.session_state.model_settings["temperature"]
        )
        
        system_prompt = f"""당신은 전문 자기소개서 작성 코치입니다.
        톤: {st.session_state.model_settings["tone"]}
        최대 길이: {st.session_state.model_settings["max_length"]}자
        
        - 구체적이고 실용적인 조언
        - 예시를 들어 설명
        - 친근하면서도 전문적인 톤
        - 이모지는 최소한으로 사용"""
        
        if uploaded_file:
            try:
                if uploaded_file.name.endswith('.txt'):
                    content = uploaded_file.read().decode('utf-8')
                elif uploaded_file.name.endswith('.docx') and DOC_LIBS_AVAILABLE:
                    doc = Document(uploaded_file)
                    content = '\n'.join([p.text for p in doc.paragraphs])
                else:
                    content = "파일을 읽을 수 없습니다."
                
                user_input = f"다음 자기소개서를 검토하고 개선점을 제안해주세요:\n\n{content}\n\n{user_input}"
            except Exception as e:
                return f"파일 처리 중 오류: {e}"
        
        prompt = ChatPromptTemplate.from_messages([
            ("system", system_prompt),
            ("human", "{input}")
        ])
        
        chain = LLMChain(llm=llm, prompt=prompt)
        response = chain.invoke({"input": user_input})
        
        return response.get("text", str(response))
        
    except Exception as e:
        return f"오류가 발생했습니다. 다시 시도해주세요.\n{str(e)}"

# ================= 대화 저장 =================
def save_conversation():
    content = ""
    for msg in st.session_state.messages:
        role = "👤 사용자" if msg["role"] == "user" else "🤖 AI 코치"
        content += f"[{msg.get('time', '')}] {role}\n{msg['content']}\n\n"
    
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    filename = f"자소서대화_{timestamp}"
    
    # 선택된 형식으로 저장
    if st.session_state.save_format == "txt":
        file_data = content
        mime = "text/plain"
        ext = "txt"
    elif st.session_state.save_format == "docx" and DOC_LIBS_AVAILABLE:
        doc = Document()
        doc.add_heading('AI 자기소개서 코칭 대화', 0)
        for para in content.split('\n'):
            doc.add_paragraph(para)
        
        bio = io.BytesIO()
        doc.save(bio)
        file_data = bio.getvalue()
        mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext = "docx"
    else:
        file_data = content
        mime = "text/plain"
        ext = "txt"
    
    # 저장 목록에 추가
    st.session_state.saved_files.append({
        "name": f"{filename}.{ext}",
        "date": datetime.datetime.now().strftime("%Y-%m-%d %H:%M"),
        "size": len(file_data),
        "data": file_data,
        "mime": mime
    })
    
    return f"{filename}.{ext}"

# ================= UI 렌더링 함수 =================
def render_header():
    st.markdown(f'''
        <div class="chat-header">
            <div class="header-title">AI 자기소개서 코칭</div>
        </div>
    ''', unsafe_allow_html=True)

def render_chat_tab():
    st.markdown("### 💬 AI 자기소개서 코칭")
    
    # 채팅 메시지 표시
    chat_container = st.container()
    with chat_container:
        st.markdown('<div class="chat-container">', unsafe_allow_html=True)
        
        for msg in st.session_state.messages:
            if msg["role"] == "user":
                st.markdown(f'''
                    <div class="msg-row user">
                        <div class="msg-time">{msg.get("time", "")}</div>
                        <div class="msg-bubble user">{msg["content"]}</div>
                    </div>
                ''', unsafe_allow_html=True)
            else:
                content_html = msg["content"].replace('\n', '<br>')
                st.markdown(f'''
                    <div class="msg-row ai">
                        <div class="msg-bubble ai">{content_html}</div>
                        <div class="msg-time">{msg.get("time", "")}</div>
                    </div>
                ''', unsafe_allow_html=True)
        
        st.markdown('</div>', unsafe_allow_html=True)
    
    # 입력 영역
    input_container = st.container()
    with input_container:
        st.markdown('<div class="input-container">', unsafe_allow_html=True)
        
        # 빠른 답변
        st.markdown("**빠른 답변:**")
        quick_replies = ["가이드라인 알려줘", "자소서 시작하기", "첨삭 받고 싶어", "예시 보여줘"]
        cols = st.columns(len(quick_replies))
        for i, reply in enumerate(quick_replies):
            with cols[i]:
                if st.button(reply, key=f"quick_{i}", use_container_width=True):
                    st.session_state.messages.append({
                        "role": "user",
                        "content": reply,
                        "time": datetime.datetime.now().strftime("%H:%M")
                    })
                    response = get_ai_response(reply)
                    st.session_state.messages.append({
                        "role": "ai",
                        "content": response,
                        "time": datetime.datetime.now().strftime("%H:%M")
                    })
                    st.rerun()
        
        # 파일 업로드
        uploaded_file = st.file_uploader(
            "📎 파일 첨부 (txt, docx)",
            type=['txt', 'docx'],
            help="자기소개서 파일을 업로드하면 첨삭을 도와드립니다"
        )
        
        # 메시지 입력
        with st.form("chat_form", clear_on_submit=True):
            col1, col2 = st.columns([5, 1])
            with col1:
                user_input = st.text_input(
                    "메시지",
                    placeholder="메시지를 입력하세요...",
                    label_visibility="collapsed"
                )
            with col2:
                send = st.form_submit_button("전송", use_container_width=True)
            
            if send and user_input:
                # 사용자 메시지 추가
                st.session_state.messages.append({
                    "role": "user",
                    "content": user_input,
                    "time": datetime.datetime.now().strftime("%H:%M")
                })
                
                # AI 응답 생성
                with st.spinner("입력 중..."):
                    response = get_ai_response(user_input, uploaded_file)
                
                st.session_state.messages.append({
                    "role": "ai",
                    "content": response,
                    "time": datetime.datetime.now().strftime("%H:%M")
                })
                
                st.rerun()
        
        st.markdown('</div>', unsafe_allow_html=True)

def render_settings_tab():
    st.markdown("### ⚙️ 설정")
    
    # API 설정
    with st.expander("🔑 API 설정", expanded=True):
        api_key = st.text_input(
            "OpenAI API Key",
            value=st.session_state.api_key,
            type="password",
            placeholder="sk-...",
            help="OpenAI API 키를 입력하세요"
        )
        
        if api_key != st.session_state.api_key:
            st.session_state.api_key = api_key
            st.success("API 키가 저장되었습니다!")
        
        st.info("💡 API 키가 없어도 기본 기능을 사용할 수 있습니다")
    
    # 대화 관리
    with st.expander("💬 대화 관리", expanded=True):
        col1, col2 = st.columns(2)
        with col1:
            if st.button("🗑️ 대화 초기화", use_container_width=True):
                st.session_state.messages = [{
                    "role": "ai",
                    "content": "안녕하세요! AI 자기소개서 코치입니다. 무엇을 도와드릴까요?",
                    "time": datetime.datetime.now().strftime("%H:%M")
                }]
                st.success("대화가 초기화되었습니다!")
                st.rerun()
        
        with col2:
            if st.button("💾 대화 저장", use_container_width=True):
                filename = save_conversation()
                st.success(f"{filename} 저장됨!")

def render_advanced_settings_tab():
    st.markdown("### 🔧 세부설정")
    
    # AI 모델 설정
    with st.expander("🤖 AI 모델 설정", expanded=True):
        st.session_state.model_settings["temperature"] = st.slider(
            "창의성 (Temperature)",
            min_value=0.0,
            max_value=1.0,
            value=st.session_state.model_settings["temperature"],
            step=0.1,
            help="높을수록 창의적이고 다양한 응답을 생성합니다"
        )
        
        st.session_state.model_settings["max_length"] = st.number_input(
            "최대 응답 길이 (자)",
            min_value=100,
            max_value=3000,
            value=st.session_state.model_settings["max_length"],
            step=100,
            help="AI 응답의 최대 길이를 설정합니다"
        )
        
        st.session_state.model_settings["tone"] = st.selectbox(
            "응답 톤",
            ["professional", "friendly", "casual", "formal"],
            index=["professional", "friendly", "casual", "formal"].index(st.session_state.model_settings["tone"]),
            help="AI 응답의 톤을 설정합니다"
        )
    
    # 저장 설정
    with st.expander("💾 저장 설정", expanded=True):
        st.session_state.save_format = st.selectbox(
            "기본 저장 형식",
            ["txt", "docx", "pdf"],
            index=["txt", "docx", "pdf"].index(st.session_state.save_format),
            help="대화 저장 시 기본 형식을 설정합니다"
        )
        
        st.info("📌 저장된 파일은 '저장소' 탭에서 확인할 수 있습니다")

def render_storage_tab():
    st.markdown("### 📁 저장소")
    
    if not st.session_state.saved_files:
        st.info("저장된 파일이 없습니다. 대화를 저장하려면 설정 탭을 이용하세요.")
    else:
        for i, file in enumerate(st.session_state.saved_files):
            col1, col2 = st.columns([4, 1])
            with col1:
                st.markdown(f'''
                    <div class="file-item">
                        <div class="file-info">
                            <div class="file-name">📄 {file["name"]}</div>
                            <div class="file-date">{file["date"]} · {file["size"]} bytes</div>
                        </div>
                    </div>
                ''', unsafe_allow_html=True)
            
            with col2:
                st.download_button(
                    label="다운로드",
                    data=file["data"],
                    file_name=file["name"],
                    mime=file["mime"],
                    key=f"download_{i}"
                )
        
        # 일괄 삭제
        if st.button("🗑️ 모든 파일 삭제", use_container_width=True):
            st.session_state.saved_files = []
            st.success("모든 파일이 삭제되었습니다!")
            st.rerun()

# ================= 메인 앱 =================
def main():
    # 헤더
    render_header()
    
    # 탭 생성
    tab1, tab2, tab3, tab4 = st.tabs(["💬 대화", "⚙️ 설정", "🔧 세부설정", "📁 저장소"])
    
    with tab1:
        render_chat_tab()
    
    with tab2:
        render_settings_tab()
    
    with tab3:
        render_advanced_settings_tab()
    
    with tab4:
        render_storage_tab()

if __name__ == "__main__":
    main()
